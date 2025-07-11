import argparse
import glob
import os
import shutil
import re
import docx2pdf
import openpyxl
import openpyxl.utils
import pandas as pd
import numpy as np
from datetime import datetime
from io import BytesIO
from openpyxl.worksheet.datavalidation import DataValidation
import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import Paragraph, SimpleDocTemplate, Image, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from collections import defaultdict
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import duckdb
import string

import mapping

DEFAULT_PATH = None      # "C:\\Users\\Pfad\\Zu\\checklisten_2023"

DEFAULT_SCALE = '"Ja, Teilweise, Nein"' # Für scale_setter() relevant, sonst nicht. Format der Anführungsstriche muss beibehalten werden.

# Gibt zurück, ob eine Datei dem Format der Checkliste entspricht
def is_format(filename):
    return filename.endswith('.xlsx') and "Checkliste" in filename

# Gibt den Namen eines Bausteins anhand des Dateinamens zurück, entweder kurz "APP.1.1" oder lang "Office Produkte"
def get_name(filename, short):
    filename_pattern = re.compile(r"(?<=\bCheckliste[_\s])([A-Z]{3,4}(?:\.[0-9]+){1,3})(?:.*?)?(?=\.xlsx$)")

    if is_format(filename):
        match = filename_pattern.search(filename)
        if match:
            ref = match.group(1)
            if short:
                return ref
            return mapping.bsi_ref_titles.get(ref, "Unbekanntes Kürzel")
        return None

    elif is_pdf(filename) or is_docx(filename):
        ref = filename.split()[0]
        if ref.startswith("CCON"):
            ref = "CON" + ref[4:]
        if short:
            return ref
        return mapping.bsi_ref_titles.get(ref, "Unbekanntes Kürzel")
    return None

# Gibt das Bausteinkürzel einer Anforderungs-ID zurück
def get_baustein_from_id(anforderung_id):
    if anforderung_id == "OPS.2.3A22":
        return "OPS.2.3"
    position = anforderung_id.find('.A')
    return anforderung_id[:position]

# Überprüft, ob eine PDF-Datei den Namen eines IT-Grundschutz-Bausteins hat
def is_pdf(file):
    if file.endswith(".pdf") or file.endswith(".pdf.clean"):
        split = file.split()
        if split:
            ref = split[0]
            if ref.startswith("CCON"):
                ref = "CON" + ref[4:]
            if ref in mapping.bsi_ref_titles:
                return True
    return False

# Überprüft, ob eine DOCX-Datei den Namen eines IT-Grundschutz-Bausteins hat
def is_docx(file):
    if file.endswith(".docx"):
        split = file.split()
        if split:
            ref = split[0]
            if ref.startswith("CCON"):
                ref = "CON" + ref[4:]
            if ref in mapping.bsi_ref_titles:
                return True
    return False

# Lädt Exceldateien und überspringt Metadaten
def load_data(file_path):
    sheet_name = get_name(file_path, True)
    try:
        return pd.read_excel(file_path, sheet_name=sheet_name, skiprows=4)
    except PermissionError:
        print(f"Fehler: Die Datei {get_name(file_path, True)} ist möglicherweise geöffnet und kann nicht gelesen werden. Vorgang wird abgebrochen.")
    except ValueError:
        print(f"Fehler: Der erwartete Tabellenblattname '{sheet_name}' wurde in der Datei '{file_path}' nicht gefunden.")
    except Exception as e:
        print(f"Ein unerwarteter Fehler ist beim Öffnen der Datei {file_path} aufgetreten: {e}")
    exit(1)

# Fragt den User eine Frage und gibt True/False zurück
def ask_user(prompt, default_yes=True):
    if default_yes:
        user_input = input(f"\n{prompt} (Ja/nein): ").strip().lower()
        return user_input != "nein" and user_input != "n"
    else:
        user_input = input(f"\n{prompt} (ja/Nein): ").strip().lower()
        return user_input == "ja" or user_input == "j"

# Gibt dem Nutzer Auswahlmöglichkeiten, entweder eine oder mehrere Antworten können ausgewählt werden
def multiple_choice(options_dict, prompt_message, multi, default_value=None):
    print(prompt_message)
    for key, value in options_dict.items():
        print(f"{key}. {value}")

    while True:
        user_input = input("\nAuswahl: ").strip()

        if user_input == "" and default_value is not None:
            return default_value if isinstance(default_value, list) else [default_value]

        if multi:
            choices = []
            try:
                selected_keys = [c.strip() for c in user_input.split(',')]
                for key in selected_keys:
                    if key in options_dict:
                        choices.append(key)
                    else:
                        print(f"'{key}' ist keine gültige Auswahl. Bitte geben Sie gültige Nummern ein.")
                        choices = []
                        break
                if choices:
                    return choices
            except Exception:
                print("Ungültige Eingabe. Bitte geben Sie kommagetrennte Nummern ein (z.B. 1,3,5).")
        else:
            if user_input in options_dict:
                return user_input
            else:
                print("Ungültige Auswahl. Bitte geben Sie eine gültige Nummer ein.")

# Erkennt anhand der Dateien eines Ordners, ob ein bestimmtes IT-Grundschutz-Profil ausgewählt wurde
def match_profile(directory):
    files_in_dir = {get_name(file, True) for file in os.listdir(directory) if is_format(file)}
    matching_profiles = []
    for profile_name, profile_files in mapping.profiles_mapping.items():
        profile_set = set(getattr(mapping, profile_files))
        if files_in_dir == profile_set:
            matching_profiles.append(mapping.profile_names[profile_name])
    return matching_profiles, files_in_dir

# Überprüft, ob IDs eindeutig sind
def id_check(file_or_dir):
    if os.path.isfile(file_or_dir):
        df = load_data(file_or_dir)
        df = df[["ID-Anforderung"]].dropna()

        if df["ID-Anforderung"].astype(str).str.match(r'.*\d$').all():
            return 1 # Default: nicht eindeutig
        else:
            if df["ID-Anforderung"].duplicated().any():
                return 3 # Gemischt
            else:
                return 2 # Nur eindeutige
    else:
        results = {}
        for file_name in os.listdir(file_or_dir):
            file_path = os.path.join(file_or_dir, file_name)
            if os.path.isfile(file_path) and is_format(file_name):
                df = load_data(file_path)
                df = df[["ID-Anforderung"]].dropna()

                if df["ID-Anforderung"].astype(str).str.match(r'.*\d$').all():
                    results[file_name] = 1
                else:
                    if df["ID-Anforderung"].duplicated().any():
                        results[file_name] = 3
                    else:
                        results[file_name] = 2

        unique_values = set(results.values())
        if len(unique_values) == 1:
            return unique_values.pop()
        else:
            status_map = {1: "Nicht eindeutige IDs", 2: "Eindeutige IDs", 3: "Gemischte IDs"}
            return {file: status_map.get(status) for file, status in results.items()}

# Entfernt Anforderungen, bei denen eine spezifische Bedingung erfüllt ist, z.B. "Titel" ist "ENTFALLEN"
def remove_specific_requirements(df, where, is_value):
    return df[df[where] != is_value]

# Gibt Dataframe zurück, bei dem die Zeile ein bestimmter Wert ist
def get_specific_df(df, column, value):
    return df[(df[column].fillna('').str.strip().str.lower() == value)]

# Gibt Dataframe zurück, bei dem zwei Zeilen entweder leer sind oder einen bestimmten Wert haben
def get_specific_df_with_two_conditions_and_emptiness(df, column, value, column2, value2):
    return df[(df[column].isna() | (df[column].fillna('').str.strip().str.lower() == value.lower())) & (df[column2].isna() | (df[column2].fillna('').str.strip().str.lower() == value2.lower()))]

# Gibt Zahlen zurück, wie viele Anforderungen entbehrlich sowie nicht, teilweise oder ganz erfüllt sind
def implementation_count(df):
    not_implemented = get_specific_df_with_two_conditions_and_emptiness(df, "Umsetzung", "nein", "Entbehrlich", "nein")
    partly_implemented = df[df["Umsetzung"].fillna("").astype(str).str.strip().str.lower() == "teilweise"]
    fully_implemented = df[df["Umsetzung"].fillna("").astype(str).str.strip().str.lower() == "ja"]
    unnecessary = df[df["Entbehrlich"].fillna("").astype(str).str.strip().str.lower() == "ja"]
    return len(not_implemented), len(partly_implemented), len(fully_implemented), len(unnecessary)

# Summiert die Kostenspalte
def cost_count(df):
    df_copy = df.copy()
    df_copy["Kostenschätzung_cleaned"] = df_copy["Kostenschätzung"].astype(str).str.replace('€', '').str.replace(',', '.').str.findall(r'\d+\.?\d*').str[0]
    df_copy["Kostenschätzung_cleaned"] = pd.to_numeric(df_copy["Kostenschätzung_cleaned"], errors='coerce')
    return df_copy["Kostenschätzung_cleaned"].fillna(0).sum()

# Zählt wie viele Basis, Standard und erhöhte Schutzbedarfsanforderungen vorliegen
def get_type(df, default_len=True):
    basis = df[df["Typ"] == "Basis"]
    standard = df[df["Typ"] == "Standard"]
    high = df[df["Typ"] == "Hoch"]
    if default_len:
        return len(basis), len(standard), len(high)
    else: return basis, standard, high

# Gibt Verantwortliche zurück
def get_responsible(df):
    return set(df["Verantwortlich"].dropna().unique())

# Gibt eine eindeutige ID für eine Reihe in der Tabelle
def get_row_key(row):
    return f"{row.get('ID-Anforderung', '')}_{row.get('Inhalt', '')}"

# Gibt eine Prozentzahl an bisher umgesetzten Anforderungen zurück
def get_percentage_missing(df):
    df = df[["Titel", "Typ", "Entbehrlich", "Umsetzung"]]
    df = remove_specific_requirements(df, "Titel", "ENTFALLEN")
    not_implemented, partly_implemented, fully_implemented, unnecessary = implementation_count(df)
    total = not_implemented + partly_implemented + fully_implemented
    return (fully_implemented / total) * 100

# Färbt Prozentzahl ein
def get_colored_percentage(percentage):
    if percentage < 50:
        color = "\033[41m"  # Roter Hintergrund
    elif percentage < 75:
        color = "\033[43m"  # Gelber Hintergrund
    else:
        color = "\033[42m"  # Grüner Hintergrund

    reset = "\033[0m"  # Reset der Farbe
    return f"{color} {percentage:.2f}% {reset}"

# Ordnet den Tabellen ihren Namen zu und gibt aus, wenn die Dateien einem IT-Grundschutz-Profil entsprechen
def map_xlsx_files(path):
    if os.path.isfile(path):
        file_name = os.path.basename(path)
        mapped_name = get_name(file_name, False)
        print(f"{file_name} → {mapped_name if mapped_name else 'Keine Zuordnung gefunden'}")
    elif os.path.isdir(path):
        for filename in os.listdir(path):
            mapped_name = get_name(filename, False)
            print(f"{filename} → {mapped_name if mapped_name else 'Keine Zuordnung gefunden'}")

        matching_profiles, files_in_dir = match_profile(path)

        if matching_profiles and len(files_in_dir) < 111:
            if len(matching_profiles) == 1:
                print(f"\nZugeordnetes IT-Grundschutz-Profil: {matching_profiles[0]}")
            else:
                print(f"\nPassende zugeordnete IT-Grundschutz-Profile: {', '.join(matching_profiles)})")

    else: print("Datei oder Ordner nicht gefunden")

# Analysiert eine Tabelle nach diversen Aspekten
def analyze_single_file(file_path):
    if not is_format(file_path):
        print("Keine Datei im Format Checkliste_XXX.X.X.xlsx")
        return
    df = load_data(file_path)

    df = df[["Titel", "Typ", "Entbehrlich", "Umsetzung", "Verantwortlich", "Kostenschätzung"]]

    df = remove_specific_requirements(df, "Titel", "ENTFALLEN")

    not_implemented, partly_implemented, fully_implemented, unnecessary = implementation_count(df)
    cost = cost_count(df)
    all_responsibles = f"{', '.join(get_responsible(df))}"

    ni = get_specific_df_with_two_conditions_and_emptiness(df, "Umsetzung", "nein", "Entbehrlich", "nein")
    ni_basis, ni_standard, ni_high = get_type(ni)

    covered_risks, covered_cia = risk_analysis(file_path)

    print(f"Gesamtanzahl Anforderungen: {len(df)}")
    print(f"Umgesetzte Anforderungen: {fully_implemented}")
    print(f"Teilweise umgesetzte Anforderungen: {partly_implemented}")
    print(f"Entbehrliche Anforderungen: {unnecessary}")
    print(f"Nicht umgesetzte Anforderungen: {not_implemented}")
    print(f"(davon Basis: {ni_basis}, Standard: {ni_standard}, Erhöhter Schutzbedarf: {ni_high})")
    print(f"Summierte Kostenschätzung: {cost}€")
    print(f"Verantwortliche: {all_responsibles}")
    print(f"(Teilweise) abgedeckte elementare Gefährdungen: {len(covered_risks)}/47")
    print(f"(Teilweise) abgedeckte Schutzziele: {len(covered_cia)}/3")

# Analysiere alle Tabellen nach diversen Aspekten
def analyze_all_files(folder_path):
    total_files = 0
    total_df_count = 0
    total_fully_implemented = 0
    total_partly_implemented = 0
    total_unnecessary = 0
    total_not_implemented = 0
    total_cost = 0
    total_basis = 0
    total_standard = 0
    total_high = 0
    total_risks = set()
    total_cia = set()

    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        if is_format(file_path):
            total_files += 1
            df = load_data(file_path)

            df = df[["Titel", "Typ", "Entbehrlich", "Umsetzung", "Kostenschätzung"]]

            df = remove_specific_requirements(df, "Titel", "ENTFALLEN")

            not_implemented, partly_implemented, fully_implemented, unnecessary = implementation_count(df)
            cost = cost_count(df)

            ni = get_specific_df_with_two_conditions_and_emptiness(df, "Umsetzung", "nein", "Entbehrlich", "nein")
            ni_basis, ni_standard, ni_high = get_type(ni)

            covered_risks, covered_cia = risk_analysis(file_path)

            total_df_count += len(df)
            total_fully_implemented += fully_implemented
            total_partly_implemented += partly_implemented
            total_unnecessary += unnecessary
            total_not_implemented += not_implemented
            total_cost += cost
            total_basis += ni_basis
            total_standard += ni_standard
            total_high += ni_high
            total_risks.update(covered_risks)
            total_cia.update(covered_cia)
    if total_files == 0:
        print("Keine Dateien im Format Checkliste_XXX.X.X.xlsx gefunden")
        return
    total_risks_num = len(total_risks)
    total_cia_num = len(total_cia)

    print(f"Analysierte Tabellen: {total_files}")
    print(f"Gesamtanzahl Anforderungen: {total_df_count}")
    print(f"Umgesetzte Anforderungen: {total_fully_implemented}")
    print(f"Teilweise umgesetzte Anforderungen: {total_partly_implemented}")
    print(f"Entbehrliche Anforderungen: {total_unnecessary}")
    print(f"Nicht umgesetzte Anforderungen: {total_not_implemented}")
    print(f"(davon Basis: {total_basis}, Standard: {total_standard}, Erhöhter Schutzbedarf: {total_high})")
    print(f"Summierte Kostenschätzung: {total_cost}€")
    print(f"(Teilweise) abgedeckte elementare Gefährdungen: {(total_risks_num)}/47")
    print(f"(Teilweise) abgedeckte Schutzziele: {total_cia_num}/3")

# Behandelt die Prozentzahl an umgesetzten Anforderungen für eine Datei oder einen Ordner
# Funktioniert nicht mit altem Windows cmd
def percentages(path):
    if os.path.isfile(path) and is_format(path):
        df = load_data(path)
        df = df[["Titel", "Typ", "Entbehrlich", "Umsetzung"]]
        df = remove_specific_requirements(df, "Titel", "ENTFALLEN")
        percentage = get_percentage_missing(df)
        print(f"Insgesamt vollständig erfüllte Anforderungen von {get_name(path, True)}: {get_colored_percentage(percentage)}")

        df_basis, df_standard, df_high = get_type(df, default_len=False)
        basis_missing = get_percentage_missing(df_basis)
        standard_missing = get_percentage_missing(df_standard)
        high_missing = get_percentage_missing(df_high)
        print(f"Basis-Anforderungen: {get_colored_percentage(basis_missing)}")
        print(f"Standard-Anforderungen: {get_colored_percentage(standard_missing)}")
        print(f"Anforderungen des erhöhten Schutzbedarfs: {get_colored_percentage(high_missing)}")

    elif os.path.isdir(path):
        print("Baustein → Prozent erfüllt")
        for file_name in os.listdir(path):
            file_path = os.path.join(path, file_name)
            if is_format(file_path):
                df = load_data(file_path)
                df = df[["Titel", "Typ", "Entbehrlich", "Umsetzung"]]
                df = remove_specific_requirements(df, "Titel", "ENTFALLEN")
                percentage = get_percentage_missing(df)
                print(f"{get_name(file_name, True)} → {get_colored_percentage(percentage)}")
    else: print("Datei oder Ordner nicht gefunden")

# Behandelt die IT-Grundschutz-Profilauswahl und führt den Verschiebungsprozess von Dateien durch
def profile_handler(file_or_dir):
    if os.path.isdir(file_or_dir):

        print("Folgende IT-Grundschutz-Profile stehen zur Verfügung (auch auf PDF, DOCX anwendbar):")
        for key, value in mapping.profile_names.items():
            print(f"{key}: {value}")
        print("Wähle mit der Zahl aus. Hinweis: Es werden nur die Bausteine, die auf jeden Fall ausgeschlossen sind, verschoben.\nManche Profile besitzen außerdem zusätzliche Bausteine, die hiermit nicht installiert werden.")
        user_input = input("\nAuswahl (oder exit): ").strip()
        if user_input.lower() == "exit":
            print("Abgebrochen")

        else:
            if user_input in mapping.profiles_mapping:
                profile_name = mapping.profiles_mapping[user_input]
                profile = getattr(mapping, profile_name)
            else:
                print("Ungültige Eingabe. Vorgang wird abgebrochen.")
                return

            target_dir = os.path.join(file_or_dir, "Nach Profil nicht benötigt")
            os.makedirs(target_dir, exist_ok=True)

            moved = False
            for file in os.listdir(file_or_dir):
                if is_format(file) or is_pdf(file) or is_docx(file):
                    file_path = os.path.join(file_or_dir, file)
                    file_short = get_name(file, True)
                    if file_short not in profile and file_short in mapping.bsi_ref_titles:
                        shutil.move(file_path, os.path.join(target_dir, file))
                        print(f"{get_name(file, True)}: {get_name(file, False)} wurde entfernt")
                        moved = True
            if not moved:
                print("Keine Datei entfernt, entweder fehlen Dateien oder das Profil ist nur ein Platzhalter und muss manuell ergänzt werden")

    else: print("Verzeichnis nicht gefunden")

# Setzt leere Standard- oder erhöhte Schutzbedarfsanforderungen auf entbehrlich
def ignore_empty_requirements(file_path, column, value, second_empty_no_type, fully):
    df = load_data(file_path)

    if not fully:
        mask = (df[column] == value) & (df["Entbehrlich"].isna()) & (df[second_empty_no_type].isna())
    else:
        mask = (df[column] == value) & (df["Entbehrlich"].isna())

    wb = openpyxl.load_workbook(file_path)
    sheet_name = get_name(file_path, True)
    ws = wb[sheet_name]

    for idx in df[mask].index:
        excel_row = idx + 6
        col_letter = openpyxl.utils.get_column_letter(df.columns.get_loc("Entbehrlich") + 1)
        ws[f"{col_letter}{excel_row}"].value = "Ja"

        bemerkung_col = openpyxl.utils.get_column_letter(df.columns.get_loc("Begründung für Entbehrlichkeit") + 1)
        if ws[f"{bemerkung_col}{excel_row}"].value:
            ws[f"{bemerkung_col}{excel_row}"].value += f"- Alle Anforderungen des Typs {value} auf Entbehrlich gesetzt"
        else:
            ws[f"{bemerkung_col}{excel_row}"].value = f"Alle Anforderungen des Typs {value} auf Entbehrlich gesetzt"

    wb.save(file_path)

# Handler für Datei oder Ordner
def ignore_empty_requirements_handler(path, column, value, second_empty_type, fully):
    if os.path.isdir(path):
        for file in os.listdir(path):
            if is_format(file):
                file_path = os.path.join(path, file)
                ignore_empty_requirements(file_path, column, value, second_empty_type, fully)
        if not fully:
            print(f"Alle leeren {value}-Anforderungen aus {path} wurden auf entbehrlich gesetzt und gespeichert.")
        else:
            print(f"Alle {value}-Anforderungen aus {path} wurden auf entbehrlich gesetzt und gespeichert.")
    elif os.path.isfile(path) and is_format(path):
        ignore_empty_requirements(path, column, value, second_empty_type, fully)
        if not fully:
            print(f"Alle leeren {value}-Anforderungen von {get_name(path, True)} wurden auf entbehrlich gesetzt und gespeichert.")
        else:
            print(f"Alle {value}-Anforderungen von {get_name(path, True)} wurden auf entbehrlich gesetzt und gespeichert.")
    else: print("Keine Datei(en) mit passendem Format gefunden")

# Sortiert Dateien nach verschiedenen Kriterien
def sort_list(directory):
    if not os.path.isdir(directory):
        print("Ordner nicht gefunden")
        return

    print("Die Dateien können nach folgenden Kriterien absteigend sortiert werden:")
    for key, value in mapping.sort_list.items():
        print(f"{key}: {value}")
    print("Wähle mit der Zahl aus.")
    sort_key = input("\nAuswahl (oder exit): ").strip()
    if sort_key.lower() == "exit":
        print("Abgebrochen")
        return

    if sort_key not in mapping.sort_list:
        print("Ungültige Auswahl")
        return

    files_data = []

    for file_name in os.listdir(directory):
        file_path = os.path.join(directory, file_name)
        if is_format(file_name):
            df = load_data(file_path)
            df = df[["Titel", "Typ", "Entbehrlich", "Umsetzung", "Umsetzung bis", "Verantwortlich", "Kostenschätzung"]]
            df = remove_specific_requirements(df, "Titel", "ENTFALLEN")

            current_file_results = {"file": file_name}

            if sort_key in ["1", "2", "3", "8", "10", "11"]:
                not_implemented, partly_implemented, fully_implemented, unnecessary = implementation_count(df)
                if sort_key == "1":
                    current_file_results[sort_key] = not_implemented + partly_implemented
                elif sort_key == "2":
                    current_file_results[sort_key] = not_implemented
                elif sort_key == "3":
                    current_file_results[sort_key] = fully_implemented
                elif sort_key == "8":
                    current_file_results[sort_key] = partly_implemented
                elif sort_key == "10":
                    current_file_results[sort_key] = unnecessary
                elif sort_key == "11":
                    current_file_results[sort_key] = len(df[df["Entbehrlich"].fillna("").astype(str).str.strip().str.lower() == "nein"])

            elif sort_key in ["4", "5", "6", "7"]:
                ni = get_specific_df_with_two_conditions_and_emptiness(df, "Umsetzung", "nein", "Entbehrlich", "nein")
                ni_basis, ni_standard, ni_high = get_type(ni)
                if sort_key == "4":
                    current_file_results[sort_key] = ni_basis
                elif sort_key == "5":
                    current_file_results[sort_key] = ni_standard
                elif sort_key == "6":
                    current_file_results[sort_key] = ni_high
                elif sort_key == "7":
                    current_file_results[sort_key] = ni_basis + ni_standard

            elif sort_key == "9":
                current_file_results[sort_key] = get_percentage_missing(df)

            elif sort_key == "12":
                current_file_results[sort_key] = cost_count(df)

            elif sort_key == "13":
                current_file_results[sort_key] = get_responsible(df)

            elif sort_key == "14":
                current_file_results[sort_key] = len(df)

            elif sort_key == "15":
                current_file_results[sort_key] = ", ".join(sorted(list(map(format_dates, df["Umsetzung bis"].dropna().unique()))))

            elif sort_key in ["16", "17"]:
                covered_risks, covered_cia = risk_analysis(file_path)
                if sort_key == "16":
                    current_file_results[sort_key] = len(covered_risks)
                elif sort_key == "17":
                    current_file_results[sort_key] = len(covered_cia)

            if sort_key in current_file_results:
                files_data.append(current_file_results)

    if not files_data:
        print("Keine passenden Dateien gefunden.")
        return

    files_data.sort(key=lambda x: x[sort_key], reverse=True)

    print(f"Dateien sortiert nach {mapping.sort_list[sort_key]}:\n")
    for item in files_data:
        print(f"{get_name(item['file'], True)}: {item[sort_key]}")

# Erkennt WiBA-Anforderungen anhand des Textes in einer Datei und setzt sie auf umgesetzt
def wiba_setter(file_path, wiba, wiba_ids):
    df = load_data(file_path)

    wiba_mapping = {w: i for i, w in enumerate(wiba)}

    id_status = id_check(file_path)
    if id_status == 1:
        ids = df["ID-Anforderung"]
    elif id_status == 2:
        ids = pd.Series([req_id[:-1] for req_id in df["ID-Anforderung"]])
    else:
        print(f"Sowohl eindeutige als auch nicht eindeutige IDs in {file_path} gefunden, eventuell können nicht alle WiBA-Anforderungen zugeordnet werden")
        ids = df["ID-Anforderung"]

    # Überprüfung, ob Inhalt und ID gleich sind, da der Inhalt einiger Anforderungen doppelt vorkommt
    mask = (
            df["Inhalt"].isin(wiba) &
            ids.isin(wiba_ids) &
            df["Inhalt"].map(wiba_mapping).notna() &
            (ids == df["Inhalt"].map(
                lambda x: wiba_ids[wiba_mapping[x]] if x in wiba_mapping else None)) &
            df["Umsetzung"].isna()
    )

    # Falls keine Änderungen abbrechen
    if not mask.any():
        return 0

    wb = openpyxl.load_workbook(file_path)
    sheet_name = get_name(file_path, True)
    ws = wb[sheet_name]

    for idx in df[mask].index:
        excel_row = idx + 6
        umsetzung_col = openpyxl.utils.get_column_letter(df.columns.get_loc("Umsetzung") + 1)
        bemerkung_col = openpyxl.utils.get_column_letter(df.columns.get_loc("Bemerkungen / Begründung für Nicht-Umsetzung") + 1)

        ws[f"{umsetzung_col}{excel_row}"].value = "Ja"
        if ws[f"{bemerkung_col}{excel_row}"].value:
            ws[f"{bemerkung_col}{excel_row}"].value += "- WiBA Integration"
        else:
            ws[f"{bemerkung_col}{excel_row}"].value = "WiBA Integration"
    wb.save(file_path)
    return 1

# Handler für Datei oder Ordner
def wiba_transfer(path):
    wiba = mapping.wiba
    wiba_ids = mapping.wiba_id
    if os.path.isfile(path) and is_format(path):
        print("Starte WiBA Integration...")
        wiba_setter(path, wiba, wiba_ids)
        print(f"WiBA erfolgreich in {get_name(path, True)} integriert")
    elif os.path.isdir(path):
        print("Starte WiBA Integration...")
        temp = 0
        for file_name in os.listdir(path):
            file_path = os.path.join(path, file_name)
            if is_format(file_name):
                temp += wiba_setter(file_path, wiba, wiba_ids)
        if temp:
            print(f"WiBA erfolgreich in {temp} Dateien integriert")
        else: print("Keine Dateien im Ordner, die auch in WiBA vorkommen")
    else: print("Datei oder Ordner nicht gefunden")

# Handler für Datei oder Ordner
def scale_setter_handler(file_or_dir):
    user_choice = ask_user("Möchten Sie eine eigene Skala angeben? Sonst wird der im Code definierte Default verwendet", default_yes=False)
    values = None
    if user_choice:
        num_values = int(input("Wie viele Werte? ").strip())
        values_list = []

        for i in range(num_values):
            value = input(f"Geben Sie Wert {i+1} ein: ").strip()
            values_list.append(f"{value}")
        values = '"' + ", ".join(values_list) + '"'
        print("Starte Modifikation...")
    else: print("Default Skala verwendet. Starte Modifikation...")

    if os.path.isfile(file_or_dir) and is_format(file_or_dir):
        scale_setter(file_or_dir, values)
        print(f"{file_or_dir} modifiziert und gespeichert!")
    elif os.path.isdir(file_or_dir):
        for file_name in os.listdir(file_or_dir):
            file_path = os.path.join(file_or_dir, file_name)
            if is_format(file_name):
                scale_setter(file_path, values)
        print(f"Alle Dateien in {file_or_dir} modifiziert und gespeichert!")
    else: print("Keine Datei gefunden")

# Ändert die Skala der Umsetzung in einer Tabelle
def scale_setter(file_path, values):
    df = load_data(file_path)
    cell_max = len(df) + 5
    if values is None:
        values = DEFAULT_SCALE

    wb = openpyxl.load_workbook(file_path)
    sheet_name = get_name(file_path, True)
    sheet = wb[sheet_name]

    new_validations = []

    # Entfernt bestehende Skala
    for dv in sheet.data_validations.dataValidation:
        if any("H" in str(rng) for rng in dv.ranges):
            continue
        new_validations.append(dv)  # Behalte die anderen Validierungen

    sheet.data_validations.dataValidation = new_validations

    rule = DataValidation(type="list", formula1=values, allow_blank=True)

    sheet.add_data_validation(rule)
    rule.add(f'H6:H{cell_max}')

    original_dir = os.path.dirname(file_path)
    output_dir = os.path.join(original_dir, "Modified")

    os.makedirs(output_dir, exist_ok=True)

    filename = f"Checkliste_{get_name(file_path, True)}_modified.xlsx"
    output_path = os.path.join(output_dir, filename)
    wb.save(output_path)

# Handler für Datei oder Ordner
def export_handler(file_or_dir):
    file_format, columns, omitted, unnecessary, implemented, baustein, gefahren, index_numbers, types_to_remove, merge, index_choice = get_export_settings()

    if os.path.isfile(file_or_dir) and is_format(file_or_dir):
        export(file_or_dir, file_format, columns, omitted, unnecessary, implemented, baustein, gefahren, index_numbers, types_to_remove, index_choice)
        print("\nDatei erfolgreich exportiert")
    elif os.path.isdir(file_or_dir):
        print("\nExportiere...")
        exported_count = 0
        new_files_list = []
        for file_name in os.listdir(file_or_dir):
            file_path = os.path.join(file_or_dir, file_name)
            if os.path.isfile(file_path) and is_format(file_name):  
                new_file_path = export(file_path, file_format, columns, omitted, unnecessary, implemented, baustein, gefahren, index_numbers, types_to_remove, index_choice)
                exported_count += 1
                new_files_list.append(new_file_path)
        if merge and exported_count > 1:
            merge_export(new_files_list, file_format, index_choice)
            print(f"\n{exported_count} Dateien erfolgreich exportiert und zusammengeführt.")
            return
        if exported_count > 0:
            print(f"\n{exported_count} Dateien erfolgreich exportiert.")
        else: print("Keine passenden Excel-Dateien im Ordner gefunden.")
    else:
        print(f"'{file_or_dir}' ist weder eine passende Excel-Datei noch ein Ordner.")

# Fügt die exportierten Dateien zusammen und löscht die Einzeldateien
def merge_export(file_list, file_format, index_choice):
    print("\nFüge Dateien zusammen...")
    dfs_to_merge = []

    for f_path in file_list:
        if file_format == "1":
            df = pd.read_excel(f_path)
        elif file_format == "2":
            df = pd.read_csv(f_path)
        elif file_format == "3":
            df = pd.read_json(f_path, orient="records")
        elif file_format in ["7", "8"] or file_format.startswith("delimited:"):
            delimiter = " " if file_format == "7" else "\t" if file_format == "8" else \
                file_format.split(":")[1]
            df = pd.read_csv(f_path, sep=delimiter)
        else:
            print("\nWarnung: Zusammenführen für das gewählte Format wird nicht unterstützt. Dieser Schritt wird übersprungen.")
            print(f"\nDateien erfolgreich exportiert.")
            exit(0)
        dfs_to_merge.append(df)

    merged_df = pd.concat(dfs_to_merge, ignore_index=True)

    export_dir = os.path.dirname(file_list[0])
    merged_file_base = os.path.join(export_dir, "Export")

    index_numbers = index_choice in ["2", "3"]

    if file_format == "1":
        merged_df.to_excel(merged_file_base + ".xlsx", index=index_numbers)
    elif file_format == "2":
        merged_df.to_csv(merged_file_base + ".csv", index=index_numbers)
    elif file_format == "3":
        merged_df.to_json(merged_file_base + ".json", orient="records", indent=4, force_ascii=False)
    elif file_format == "7":
        merged_df.to_csv(merged_file_base + ".txt", sep=" ", index=index_numbers)
    elif file_format == "8":
        merged_df.to_csv(merged_file_base + ".tsv", sep="\t", index=index_numbers)
    elif file_format.startswith("delimited:"):
        delimiter = file_format.split(":")[1]
        merged_df.to_csv(merged_file_base + ".txt", sep=delimiter, index=index_numbers)

    for f_path in file_list:
        os.remove(f_path)

# Fragt den Nutzer, welches Format und welche Spalten er exportieren möchte
def get_export_settings():
    export_format = multiple_choice(mapping.format_options, "Wählen Sie das Exportformat:", multi=False)

    if export_format == "9":
        delimiter = input("Geben Sie das Trennzeichen ein: ")
        export_format = f"delimited:{delimiter}"

    columns = [
        "ID-Anforderung", "Titel", "Inhalt", "Typ", "Entbehrlich",
        "Begründung für Entbehrlichkeit", "Umsetzung", "Umsetzung bis",
        "Verantwortlich", "Bemerkungen / Begründung für Nicht-Umsetzung", "Kostenschätzung"
    ]
    column_options = {str(i + 1): col for i, col in enumerate(columns)}

    print("\nHinweis: Falls Daten später wieder importiert werden sollen, werden ID-Anforderung und Inhalt benötigt")
    selected_column_keys = multiple_choice(column_options,"\nSpaltenauswahl (z.B. 1,3,5 oder Enter für Alle)", multi=True, default_value="Alle")

    if "Alle" in selected_column_keys:
        selected_columns = columns
    else:
        selected_columns = [columns[int(key) - 1] for key in selected_column_keys]

    type_options = {
        "1": "Basis",
        "2": "Standard",
        "3": "Hoch"
    }

    types_to_remove_keys = multiple_choice(type_options,
                                           "\nAuswahl zur Entfernung der Schutzbedarfstypen (z.B. 2,3 oder Enter für Keine):",
                                           multi=True, default_value=[])

    types_to_remove = [type_options[key] for key in types_to_remove_keys if key in type_options]

    omitted = ask_user("Möchten Sie entfallene Anforderungen entfernen?")
    unnecessary = ask_user("Möchten Sie als entbehrlich markierte Anforderungen entfernen?", default_yes=False)
    implemented = ask_user("Möchten Sie bereits umgesetzte Anforderungen entfernen?", default_yes=False)
    index_numbers = False
    if export_format != "3":
        index_numbers = ask_user("Möchten Sie nummerierte Zeilen?")
    baustein = ask_user("Soll jede Anforderung den Bausteinnamen enthalten?", default_yes=False)
    gefahren = ask_user("Sollen die elementaren Gefährdungen aus den Kreuzreferenztabellen hinzugefügt werden?", default_yes=False)
    merge = ask_user("Sollen die Dateien am Ende zu einer einzelnen, großen Datei zusammengeführt werden?", default_yes=False)
    index_choice = None
    if merge and index_numbers:
        index_options = {
            "1": "Indexzahlen für jeden Baustein einzeln",
            "2": "Indexzahlen für alle Anforderungen insgesamt",
            "3": "Beides"
        }
        index_choice = multiple_choice(index_options, "\nSpezifizierung für die Indexzahlen nötig:", multi=False)

    if baustein:
        selected_columns.append("Baustein")
    if gefahren:
        selected_columns.append("Abgedeckte elementare Gefährdungen")
        selected_columns.append("Abgedeckte Schutzziele")

    return export_format, selected_columns, omitted, unnecessary, implemented, baustein, gefahren, index_numbers, types_to_remove, merge, index_choice

# Exportiert den Tabelleninhalt in ein gewünschtes Format mit diversen Anpassungsmöglichkeiten
def export(file_path, file_format, columns, omitted, unnecessary, implemented, baustein, gefahren, index_numbers, types_to_remove, index_choice):
    df = load_data(file_path)
    if omitted:
        df = remove_specific_requirements(df, "Titel", "ENTFALLEN")
    if unnecessary:
        df = remove_specific_requirements(df, "Entbehrlich", "Ja")
    if implemented:
        df = remove_specific_requirements(df, "Umsetzung", "Ja")
    if baustein:
        baustein_name = f"{get_name(file_path, True)} - {get_name(file_path, False)}"
        df["Baustein"] = baustein_name
    if gefahren:
        temp_id_anforderung = df["ID-Anforderung"].astype(str).copy()
        if id_check(file_path) == 2:
            temp_id_anforderung = temp_id_anforderung.apply(lambda x: x[:-1] if len(x) > 1 else x)
        krt_map = {entry['id']: entry for entry in mapping.krt}

        # Hilfsfunktion für effiziente Suche der Kreuzreferenztabellen
        def get_krt_eintrag(req_id):
            return krt_map.get(str(req_id), {})

        # Hilfsfunktion, um elementare Gefährdungen zuzuordnen
        def get_gefahren(req_id):
            krt_entry = get_krt_eintrag(req_id)
            gefahren_ids = krt_entry.get('gefahren')
            if gefahren_ids:
                gefahren_definitions = []
                for g_id in gefahren_ids:
                    definition = mapping.gefahren.get(g_id)
                    gefahren_definitions.append(f"{g_id}: {definition}")
                return ", ".join(gefahren_definitions)
            return ""

        # Hilfsfunktion, um Schutzziele zuzuordnen
        def get_cia(req_id):
            krt_entry = get_krt_eintrag(req_id)
            cia_string = krt_entry.get('cia')
            if cia_string:
                cia_definitions = []
                for char in cia_string:
                    definition = mapping.cia.get(char)
                    cia_definitions.append(definition)
                return ", ".join(cia_definitions)
            return "Keine"

        df["Abgedeckte elementare Gefährdungen"] = temp_id_anforderung.apply(get_gefahren)
        df["Abgedeckte Schutzziele"] = temp_id_anforderung.apply(get_cia)

    if types_to_remove:
        for t in types_to_remove:
            df = remove_specific_requirements(df, "Typ", t)
    df = df[columns]

    export_dir = os.path.join(os.path.dirname(file_path), "Export")
    os.makedirs(export_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(file_path))[0]
    export_file_base = os.path.join(export_dir, base_name)

    final_export_path = None

    if index_numbers and index_choice == "2":
        index_numbers = False

    if file_format == "1":
        final_export_path = export_file_base + ".xlsx"
        df.to_excel(final_export_path, index=index_numbers)
    elif file_format == "2":
        final_export_path = export_file_base + ".csv"
        df.to_csv(final_export_path, index=index_numbers)
    elif file_format == "3":
        final_export_path = export_file_base + ".json"
        df.to_json(final_export_path, orient="records", indent=4, force_ascii=False)
    elif file_format == "4":
        final_export_path = export_file_base + ".md"
        df.to_markdown(final_export_path, index=index_numbers)
    elif file_format == "5":
        final_export_path = export_file_base + ".html"
        df.to_html(final_export_path, index=index_numbers)
    elif file_format == "6":
        final_export_path = export_file_base + ".xml"
        df.to_xml(final_export_path, index=index_numbers)
    elif file_format == "7":
        final_export_path = export_file_base + ".txt"
        df.to_csv(final_export_path, sep=" ", index=index_numbers)
    elif file_format == "8":
        final_export_path = export_file_base + ".tsv"
        df.to_csv(final_export_path, sep="\t", index=index_numbers)
    elif file_format.startswith("delimited:"):
        delimiter = file_format.split(":")[1]
        final_export_path = export_file_base + ".txt"
        df.to_csv(final_export_path, sep=delimiter, index=index_numbers)

    return final_export_path

# Überprüft, ob es nicht umgesetzte Basis-Anforderungen gibt
def basis_risk_check(file_path):
    df = load_data(file_path)
    df = df[["ID-Anforderung", "Titel", "Inhalt", "Typ", "Entbehrlich", "Umsetzung"]]
    df = remove_specific_requirements(df, "Titel", "ENTFALLEN")

    ni = get_specific_df_with_two_conditions_and_emptiness(df, "Umsetzung", "nein", "Entbehrlich", "nein")
    ni_basis, ni_standard, ni_high = get_type(ni)
    if ni_basis != 0:
        return True
    else: return False

# Führt die erste Analyse für den Report durch
def report_analysis(file_path):
    df = load_data(file_path)
    df = df[["ID-Anforderung", "Titel", "Inhalt", "Typ", "Entbehrlich", "Begründung für Entbehrlichkeit", "Umsetzung", "Umsetzung bis", "Verantwortlich", "Bemerkungen / Begründung für Nicht-Umsetzung", "Kostenschätzung"]]
    df = remove_specific_requirements(df, "Titel", "ENTFALLEN")

    not_implemented, partly_implemented, fully_implemented, unnecessary = implementation_count(df)
    cost = cost_count(df)
    due_date = ", ".join(sorted(list(map(format_dates, df["Umsetzung bis"].dropna().unique()))))
    all_responsibles = f"{', '.join(get_responsible(df))}"
    ni = get_specific_df_with_two_conditions_and_emptiness(df, "Umsetzung", "nein", "Entbehrlich", "nein")
    ni_basis, ni_standard, ni_high = get_type(ni)

    return {
        "Gesamtanzahl Anforderungen": len(df),
        "Umgesetzte Anforderungen": fully_implemented,
        "Teilweise umgesetzte Anforderungen": partly_implemented,
        "Entbehrliche Anforderungen": unnecessary,
        "Nicht umgesetzte Anforderungen": not_implemented,
        "Nicht umgesetzte Anforderungen nach Typ": {"Basis": ni_basis, "Standard": ni_standard, "Erhöhter Schutzbedarf": ni_high},
        "Summierte Kostenschätzung": cost,
        "Verantwortliche": all_responsibles,
        "Deadlines einzelner Anforderungen": due_date,
    }

# Handler um für Dateien oder alle Dateien eines Ordners Reports zu erstellen
def report(file_or_dir):
    if os.path.isfile(file_or_dir):
        if is_format(file_or_dir):
            reports_dir = os.path.join(os.path.dirname(os.path.abspath(file_or_dir)), "Reports")
            os.makedirs(reports_dir, exist_ok=True)
            results = report_analysis(file_or_dir)
            risk = basis_risk_check(file_or_dir)
            save_results_to_pdf(results, os.path.basename(file_or_dir), reports_dir, risk, file_or_dir)
            print("Report erstellt und gespeichert.")
        else:
            print("Keine Datei im Format Checkliste_XXX.X.X.xlsx")
    elif os.path.isdir(file_or_dir):
        reports_dir = os.path.join(os.path.abspath(file_or_dir), "Reports")
        os.makedirs(reports_dir, exist_ok=True)

        choice = input("Welche Reports sollen erstellt werden?\n\n1 - Für jede Datei einen Report\n2 - Einen Gesamtreport für alle Dateien\n3 - Beides\n\nAuswahl (1-3): ")
        if choice == "2":
            whole_report(file_or_dir, reports_dir)
            return
        if choice == "3":
            whole_report(file_or_dir, reports_dir)
        if choice not in ["1", "2", "3"]:
            print("Keine gültige Auswahl, es werden Berichte für jede Datei einzeln erstellt.")

        all_results = {}
        all_risks = {}
        print("Starte Analyse der Dateien für Einzelreports...")
        for file_name in os.listdir(file_or_dir):
            file_path = os.path.join(file_or_dir, file_name)
            if is_format(file_name):
                results = report_analysis(file_path)
                risk = basis_risk_check(file_path)
                all_results[file_name] = results
                all_risks[file_name] = risk
        if not all_results:
            print("Keine gültigen Dateien gefunden.")
            return
        print("Analyse abgeschlossen\nErstelle Reports...")
        for file_name, results in all_results.items():
            save_results_to_pdf(results, file_name, reports_dir, all_risks[file_name], os.path.join(file_or_dir, file_name))
        print("Reports erstellt und gespeichert.")
    else:
        print(f"Ungültiger Pfad: {file_or_dir}")

# Baut die PDF zusammen, führt dabei Berechnungen durch um festzustellen, ob bestimmte Teile benötigt werden
def save_results_to_pdf(results, file_name, reports_dir, risk, file_path):
    df = load_data(file_path)
    pdf_path = os.path.join(reports_dir, f"Report_{get_name(file_name, True)}.pdf")
    doc = SimpleDocTemplate(pdf_path, pagesize=A4, title=f"Report für {get_name(file_name, True)} - {get_name(file_name, False)}")
    story = []
    image_path2 = None

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    story.append(Paragraph(f"Analysebericht für {file_name} - {get_name(file_name, False)}", title_style))

    normal_style = styles["Normal"]
    for key, value in results.items():
        if isinstance(value, dict):
            story.append(Paragraph(f"<b>{key}:</b>", normal_style))
            for sub_key, sub_value in value.items():
                story.append(Paragraph(f"    {sub_key}: {sub_value}", normal_style))
        else:
            story.append(Paragraph(f"<b>{key}:</b> {value}", normal_style))

    if risk:
        story.append(Paragraph("<font color=red>WARNUNG: Es sind nicht alle Basis-Anforderungen umgesetzt!</font>", normal_style))
    id_status = id_check(file_path)
    if id_status == 2:
        df["ID-Anforderung"] = df["ID-Anforderung"].astype(str).apply(lambda x: x[:-1] if len(x) > 1 else x)
        story.append(Paragraph("<font color=red>Hinweis: Die IDs der Anforderungen wurden modifiziert, sodass sie für alle Teil-Anforderungen einzigartig sind</font>", normal_style))
    if id_status == 3:
        story.append(Paragraph("<font color=red>Hinweis: Die IDs der Anforderungen wurden teilweise modifiziert, möglicherweise werden sie nicht richtig gruppiert</font>", normal_style))


    image_path1 = plot_pie_chart_to_image({
        "Umgesetzt": results["Umgesetzte Anforderungen"],
        "Teilweise umgesetzt": results["Teilweise umgesetzte Anforderungen"],
        "Nicht umgesetzt": results["Nicht umgesetzte Anforderungen"],
        "Entbehrlich": results["Entbehrliche Anforderungen"]
    }, "Umsetzungsstatus")
    story.append(Image(image_path1, width=400, height=300))

    ni_df = df
    ni_df = ni_df[["ID-Anforderung", "Titel", "Inhalt", "Typ", "Umsetzung", "Entbehrlich"]]
    ni_df = remove_specific_requirements(ni_df, "Titel", "ENTFALLEN")
    ni_df = get_specific_df_with_two_conditions_and_emptiness(ni_df, "Umsetzung", "nein", "Entbehrlich", "nein")
    if not ni_df.empty:
        story.append(Paragraph(f"Nicht umgesetzte Anforderungen", title_style))
        grouped_data = defaultdict(list)
        for _, row in ni_df.iterrows():
            grouped_data[row["ID-Anforderung"]].append(row)

        table_data = [["ID-Anforderung", "Anzahl der Teil-Anforderungen", "Titel", "Typ"]]
        for id_anforderung, rows in grouped_data.items():
            first_row = rows[0]
            table_data.append([id_anforderung, len(rows), first_row["Titel"], first_row["Typ"]])

        table = get_custom_table(table_data)
        story.append(table)

        story.append(Spacer(1, 20))
        image_path2 = plot_pie_chart_to_image(results["Nicht umgesetzte Anforderungen nach Typ"],
                                          "Nicht umgesetzte Anforderungen nach Typ aufgeschlüsselt")
        story.append(Image(image_path2, width=400, height=300))

    if results["Deadlines einzelner Anforderungen"]:
        dl_df = df
        dl_df = remove_specific_requirements(dl_df, "Titel", "ENTFALLEN")
        dl_df = remove_specific_requirements(dl_df, "Umsetzung", "Ja")
        dl_df = dl_df[["ID-Anforderung", "Titel", "Typ", "Umsetzung bis"]]
        dl_df = dl_df.dropna(subset=["Umsetzung bis"])

        story.append(Paragraph(f"Offene Deadlines", title_style))
        table_data = [["ID-Anforderung", "Titel", "Typ", "Umsetzung bis"]]
        for _, row in dl_df.iterrows():
            formatted_date = format_dates(row["Umsetzung bis"])
            table_data.append([row["ID-Anforderung"], row["Titel"], row["Typ"], formatted_date])

        table = get_custom_table(table_data)
        story.append(table)

    resp_df = df
    resp_df = remove_specific_requirements(resp_df, "Verantwortlich", "ENTFALLEN")
    resp_df = resp_df[["ID-Anforderung", "Verantwortlich"]]
    resp_df = resp_df.dropna(subset=["Verantwortlich"])
    if not resp_df.empty:
        story.append(Spacer(1, 5))
        story.append(Paragraph(f"Verantwortliche für Anforderungen", title_style))
        unique_verantwortliche = resp_df["Verantwortlich"].unique()
        table_data = [["Verantwortlich", "ID-Anforderungen"]]
        for verantwortlicher in unique_verantwortliche:
                id_anforderungen = ", ".join(resp_df[resp_df["Verantwortlich"] == verantwortlicher]["ID-Anforderung"].tolist())
                table_data.append([verantwortlicher, id_anforderungen])

        table = get_custom_table(table_data)
        story.append(table)

    cost_df = df
    cost_df = remove_specific_requirements(cost_df, "ID-Anforderung", "ENTFALLEN")
    cost_df = cost_df[["ID-Anforderung", "Kostenschätzung"]]
    cost_df = cost_df.dropna(subset=["Kostenschätzung"])

    cost_df["Kostenschätzung"] = cost_df["Kostenschätzung"].astype(str) \
        .str.replace('€', '') \
        .str.replace(',', '.') \
        .str.findall(r'\d+\.?\d*').str[0]

    cost_df["Kostenschätzung"] = pd.to_numeric(cost_df["Kostenschätzung"], errors='coerce')

    cost_df = cost_df.dropna(subset=["Kostenschätzung"])

    if not cost_df.empty:
        story.append(Spacer(1, 10))
        grouped_data = cost_df.groupby("ID-Anforderung")["Kostenschätzung"].sum().reset_index()

        plt.figure(figsize=(10, 6))
        plt.bar(grouped_data["ID-Anforderung"], grouped_data["Kostenschätzung"])
        plt.xlabel("ID der Anforderung")
        plt.ylabel("Geschätzte Kosten")
        plt.title("Kostenschätzung pro Anforderung (Summe aus Teil-Anforderungen)")
        plt.xticks(rotation=45, ha="right")
        plt.tight_layout()

        img_data = BytesIO()
        plt.savefig(img_data, format="png")
        img_data.seek(0)
        story.append(Image(img_data, width=400, height=300))
        plt.close()

    implemented = get_specific_df(df, "Umsetzung", "ja")
    grouped_items = {}

    ids = implemented["ID-Anforderung"].unique()

    for req_id in ids:
        for item in mapping.krt:
            if item['id'] == req_id:
                for risk_krz in item.get('gefahren', []):
                    if req_id not in grouped_items:
                        grouped_items[req_id] = []
                    grouped_items[req_id].append(risk_krz)
                break

    if grouped_items:
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"(Teilweise) abgedeckte elementare Gefährdungen:", title_style))
        table_data = [["ID-Anforderung", "Elementare Gefährdungen"]]

        for req_id, risks_krz_list in grouped_items.items():
            risk_names = []
            for risk_krz in risks_krz_list:
                risk_names.append(f"{risk_krz}: {mapping.gefahren[risk_krz]}")

            risks_combined = "<br/>".join(risk_names)
            table_data.append([req_id, risks_combined])

        table = get_custom_table(table_data)
        story.append(table)

    entb_df = df
    entb_df = remove_specific_requirements(entb_df, "Titel", "ENTFALLEN")
    entb_df = entb_df[["ID-Anforderung", "Entbehrlich", "Begründung für Entbehrlichkeit"]]
    entb_df = entb_df[(entb_df["Entbehrlich"] == "Ja") & (entb_df["Begründung für Entbehrlichkeit"].isna())]
    if not entb_df.empty:
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"Fehlende Begründung für entbehrliche Anforderungen:", title_style))
        bullet_list = []
        for _, row in entb_df.iterrows():
            bullet_list.append(Paragraph(f"- {row['ID-Anforderung']}", normal_style))
        story.extend(bullet_list)

    ums_df = df
    ums_df = remove_specific_requirements(ums_df, "Titel", "ENTFALLEN")
    ums_df = remove_specific_requirements(ums_df, "Entbehrlich", "Ja")
    ums_df = ums_df[["ID-Anforderung", "Umsetzung", "Bemerkungen / Begründung für Nicht-Umsetzung"]]
    ums_df = ums_df[(ums_df["Umsetzung"] == "Nein") & (ums_df["Bemerkungen / Begründung für Nicht-Umsetzung"].isna())]
    if not ums_df.empty:
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"Fehlende Begründung für nicht umgesetzte Anforderungen:", title_style))
        bullet_list = []
        for _, row in ums_df.iterrows():
            bullet_list.append(Paragraph(f"- {row['ID-Anforderung']}", normal_style))
        story.extend(bullet_list)

    doc.build(story)

    os.remove(image_path1) # Temporäre Bilder entfernen, weil es früher nicht möglich ist
    if not ni_df.empty:
        os.remove(image_path2)

# Definiert wie Tabellen im Report aussehen sollen und ermöglicht Zeilenumbrüche
def get_custom_table(table_data):
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]

    paragraph_table_data = []
    for row in table_data:
        paragraph_row = []
        for cell in row:
            paragraph_row.append(Paragraph(str(cell), normal_style))
        paragraph_table_data.append(paragraph_row)

    table = Table(paragraph_table_data)
    table.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.whitesmoke),
                               ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                               ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                               ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                               ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                               ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                               ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                               ('GRID', (0, 0), (-1, -1), 1, colors.black)]))
    return table

# Vereinheitlicht die Daten in der "Umsetzung bis" Spalte
def format_dates(date_value):
    if isinstance(date_value, np.datetime64):
        return pd.to_datetime(date_value).strftime('%d.%m.%Y')
    elif isinstance(date_value, datetime):
        return date_value.strftime('%d.%m.%Y')
    else:
        return date_value

# Erstellt ein Kreisdiagramm und speichert es temporär
def plot_pie_chart_to_image(data, title):
    labels = []
    values = []
    for label, value in data.items():
        if value > 0:
            labels.append(label)
            values.append(value)

    plt.figure(figsize=(8, 6))
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=140)
    plt.title(title)
    plt.axis('equal')

    temp_image_path = f"{title}.png"
    plt.savefig(temp_image_path)
    plt.close()
    return temp_image_path

# Erstellt einen Gesamtreport für alle Dateien
def whole_report(directory, output_directory):
    print("Starte Analysen für Gesamtreport...")
    total_files = 0
    total_df_count = 0
    total_fully_implemented = 0
    total_partly_implemented = 0
    total_unnecessary = 0
    total_not_implemented = 0
    total_cost = 0
    total_basis = 0
    total_standard = 0
    total_high = 0
    for file_name in os.listdir(directory):
        file_path = os.path.join(directory, file_name)
        if is_format(file_path):
            total_files += 1
            df = load_data(file_path)
            df = df[["Titel", "Typ", "Entbehrlich", "Umsetzung", "Kostenschätzung"]]
            df = remove_specific_requirements(df, "Titel", "ENTFALLEN")

            not_implemented, partly_implemented, fully_implemented, unnecessary = implementation_count(df)
            cost = cost_count(df)

            ni = get_specific_df_with_two_conditions_and_emptiness(df, "Umsetzung", "nein", "Entbehrlich", "nein")
            ni_basis, ni_standard, ni_high = get_type(ni)
            covered_risks, covered_cia = risk_analysis(file_path)

            total_df_count += len(df)
            total_fully_implemented += fully_implemented
            total_partly_implemented += partly_implemented
            total_unnecessary += unnecessary
            total_not_implemented += not_implemented
            total_cost += cost
            total_basis += ni_basis
            total_standard += ni_standard
            total_high += ni_high
    if total_files == 0:
        print("Keine Dateien im Format Checkliste_XXX.X.X.xlsx gefunden")
        return

    results = {
    "Analysierte Tabellen": total_files,
    "Gesamtanzahl Anforderungen": total_df_count,
    "Umgesetzte Anforderungen": total_fully_implemented,
    "Teilweise umgesetzte Anforderungen": total_partly_implemented,
    "Entbehrliche Anforderungen": total_unnecessary,
    "Nicht umgesetzte Anforderungen": total_not_implemented,
    "Nicht umgesetzte Anforderungen nach Typ": {"Basis": total_basis, "Standard": total_standard, "Erhöhter Schutzbedarf": total_high},
    "Summierte Kostenschätzung": total_cost,
    }

    print("Erstelle Tabellen und Grafiken...")

    save_whole_results_to_pdf(directory, results, output_directory)

    print("Gesamtreport erstellt und gespeichert")

# Erstellt Tabellen, Grafiken und speichert den Gesamtreport
def save_whole_results_to_pdf(directory, results, reports_dir):
    pdf_path = os.path.join(reports_dir, f"Gesamtreport_{datetime.today().strftime('%d-%m-%Y')}.pdf")
    doc = SimpleDocTemplate(pdf_path, pagesize=A4,
                            title=f"IT-Grundschutz-Report")
    story = []

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    story.append(Paragraph(f"Analysebericht vom {datetime.today().strftime('%d-%m-%Y')}", title_style))

    normal_style = styles["Normal"]
    for key, value in results.items():
        if isinstance(value, dict):
            story.append(Paragraph(f"<b>{key}:</b>", normal_style))
            for sub_key, sub_value in value.items():
                story.append(Paragraph(f"    {sub_key}: {sub_value}", normal_style))
        else:
            story.append(Paragraph(f"<b>{key}:</b> {value}", normal_style))

    matching_profiles, files_in_dir = match_profile(directory)

    if matching_profiles and len(files_in_dir)<111:
        if len(matching_profiles) == 1:
            story.append(Paragraph(f"<b>IT-Grundschutz-Profil: {matching_profiles[0]}</b>", normal_style))
        else:
            story.append(Paragraph(f"<b>Passende IT-Grundschutz-Profile: {', '.join(matching_profiles)})</b>", normal_style))

    risk_files = []
    not_implemented = {}
    ni_types ={}
    deadlines = {}
    responsible = {}
    cost = {}
    missing_reason_entb = {}
    missing_reason_ni = {}
    all_covered_risks = set()
    all_covered_cia = set()
    for file_name in os.listdir(directory):
        file_path = os.path.join(directory, file_name)
        if is_format(file_path):
            df = load_data(file_path)
            df = df[["ID-Anforderung", "Titel", "Inhalt", "Typ", "Umsetzung", "Umsetzung bis", "Bemerkungen / Begründung für Nicht-Umsetzung", "Entbehrlich", "Begründung für Entbehrlichkeit", "Verantwortlich", "Kostenschätzung"]]
            df = remove_specific_requirements(df, "Titel", "ENTFALLEN")
            if basis_risk_check(file_path):
                risk_files.append(get_name(file_path, True))
            ni_df = get_specific_df_with_two_conditions_and_emptiness(df, "Umsetzung", "nein", "Entbehrlich", "nein")
            if not ni_df.empty:
                not_implemented[get_name(file_name, True)] = len(ni_df)
                ni_types[get_name(file_name, True)] = get_type(ni_df)
            dl_df = df.dropna(subset=["Umsetzung bis"])
            if not dl_df.empty:
                deadlines[get_name(file_name, True)] = []
                for _, row in dl_df.iterrows():
                    deadlines[get_name(file_name, True)].append(format_dates(row["Umsetzung bis"]))
            resp_df = df.dropna(subset=["Verantwortlich"])
            if not resp_df.empty:
                unique_verantwortliche = resp_df["Verantwortlich"].unique()
                responsible[get_name(file_name, True)] = unique_verantwortliche
            cost_df = df.copy()
            cost_df = cost_df.dropna(subset=["Kostenschätzung"])
            if not cost_df.empty:
                cost_sum = cost_count(cost_df)
                if cost_sum != 0:
                    cost[get_name(file_name, True)] = cost_sum
            entb_df = df[(df["Entbehrlich"] == "Ja") & (df["Begründung für Entbehrlichkeit"].isna())]
            if not entb_df.empty:
                missing_reason_entb[get_name(file_name, True)] = len(entb_df)
            ums_df = remove_specific_requirements(df, "Entbehrlich", "Ja")
            ums_df = ums_df[(ums_df["Umsetzung"] == "Nein") & (ums_df["Bemerkungen / Begründung für Nicht-Umsetzung"].isna())]
            if not ums_df.empty:
                missing_reason_ni[get_name(file_name, True)] = len(ums_df)
            covered_risks, covered_cia = risk_analysis(file_path)
            all_covered_risks.update(covered_risks)
            all_covered_cia.update(covered_cia)
        else: print(f"Datei/Ordner ausgelassen: {file_name}")

    print("Analysen abgeschlossen")

    if risk_files:
        story.append(Paragraph(f"<font color=red>WARNUNG: In den folgenden Bausteinen sind nicht alle Basis-Anforderungen umgesetzt: {', '.join(risk_files)}</font>",
                               normal_style))

    image_path1 = plot_pie_chart_to_image({
        "Umgesetzt": results["Umgesetzte Anforderungen"],
        "Teilweise umgesetzt": results["Teilweise umgesetzte Anforderungen"],
        "Nicht umgesetzt": results["Nicht umgesetzte Anforderungen"],
        "Entbehrlich": results["Entbehrliche Anforderungen"]
    }, "Umsetzungsstatus")
    story.append(Image(image_path1, width=400, height=300))

    story.append(Spacer(1, 20))

    image_path2 = plot_pie_chart_to_image(results["Nicht umgesetzte Anforderungen nach Typ"],
                                              "Insgesamt nicht umgesetzte Anforderungen nach Typ aufgeschlüsselt")
    story.append(Image(image_path2, width=400, height=300))

    if not_implemented:
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Nicht umgesetzte Anforderungen", title_style))
        table_data = [["Baustein", "Anzahl", "Basis", "Standard", "Erhöhter Schutzbedarf"]]
        for name, count in not_implemented.items():
            basis, standard, high = ni_types[name]
            table_data.append([name, count, basis, standard, high])
        table = get_custom_table(table_data)
        story.append(table)

    if deadlines:
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Deadlines", title_style))
        table_data = [["Baustein", "Deadlines"]]
        for name, deadlines in deadlines.items():
            deadline_str =  ", ".join(deadlines)
            table_data.append([name, deadline_str])
        table = get_custom_table(table_data)
        story.append(table)

    if responsible:
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Verantwortliche", title_style))
        table_data = [["Baustein", "Verantwortliche"]]
        for name, verantwortliche in responsible.items():
            verantwortliche_str = ", ".join(verantwortliche)
            table_data.append([name, verantwortliche_str])
        table = get_custom_table(table_data)
        story.append(table)

    if cost:
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Kostenschätzung", title_style))
        names = list(cost.keys())
        costs = list(cost.values())
        plt.figure(figsize=(10, 6))
        plt.bar(names, costs)
        plt.xlabel("Baustein")
        plt.ylabel("Kostenschätzung")
        plt.title("Summierte Kostenschätzung")
        plt.xticks(rotation=45, ha="right")
        plt.tight_layout()

        img_data = BytesIO()
        plt.savefig(img_data, format="png")
        img_data.seek(0)
        story.append(Image(img_data, width=400, height=300))
        plt.close()

    if missing_reason_entb:
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Fehlende Begründungen für Entbehrlichkeit", title_style))
        table_data = [["Baustein", "Anzahl fehlender Begründungen"]]
        for name, count in missing_reason_entb.items():
            table_data.append([name, count])
        table = get_custom_table(table_data)
        story.append(table)

    if missing_reason_ni:
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Fehlende Begründungen für Nicht-Umsetzung", title_style))
        table_data = [["Baustein", "Anzahl fehlender Begründungen"]]
        for name, count in missing_reason_ni.items():
            table_data.append([name, count])
        table = get_custom_table(table_data)
        story.append(table)

    if all_covered_risks:
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"(Teilweise) abgedeckte elementare Gefährdungen:", title_style))
        unique_risks = sorted(list(all_covered_risks))
        for risk_id in unique_risks:
            story.append(Paragraph(f"- {risk_id}: {mapping.gefahren[risk_id]}", normal_style))

    all_possible_risks = set(mapping.gefahren.keys())
    uncovered_risks = sorted(list(all_possible_risks - all_covered_risks))

    if uncovered_risks:
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"Nicht abgedeckte elementare Gefährdungen:", title_style))
        for risk_id in uncovered_risks:
            story.append(Paragraph(f"- {risk_id}: {mapping.gefahren[risk_id]}", normal_style))

    if all_covered_cia:
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"(Teilweise) abgedeckte Schutzziele:", title_style))
        unique_cia = sorted(list(all_covered_cia))
        for goal_id in unique_cia:
            story.append(Paragraph(f"- {mapping.cia[goal_id]} ({goal_id})", normal_style))

    all_possible_cia = set(mapping.cia.keys())
    uncovered_cia = sorted(list(all_possible_cia - all_covered_cia))
    if uncovered_cia:
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"Nicht abgedeckte Schutzziele:", title_style))
        for goal_id in uncovered_cia:
            story.append(Paragraph(f"- {mapping.cia[goal_id]} ({goal_id})", normal_style))

    print("Tabellen und Grafiken erstellt")

    doc.build(story)

    os.remove(image_path1)
    os.remove(image_path2)

# Löst Konflikte für verschiedene Werte in "Entbehrlich", "Umsetzung", "Umsetzung bis"
def resolve_conflict(values, column_name, file1_path, file2_path):
    unique_values = values.unique()
    if len(unique_values) == 1:
        return unique_values[0]
    elif pd.isna(unique_values[0]):
        return unique_values[1]
    elif pd.isna(unique_values[1]):
        return unique_values[0]
    else:
        if column_name == "Umsetzung bis":
            unique_values = [format_dates(val) for val in unique_values]
        conflict_row_index = values.index[0]
        conflict_row_number = conflict_row_index + 6

        print(f"\nKonflikt in Spalte '{column_name}' in Zeile {conflict_row_number}:\n")

        wb1 = openpyxl.load_workbook(file1_path)
        ws1 = wb1.active

        af_id = ws1.cell(row=conflict_row_number, column=2).value
        titel = ws1.cell(row=conflict_row_number, column=3).value
        inhalt = ws1.cell(row=conflict_row_number, column=4).value
        typ = ws1.cell(row=conflict_row_number, column=5).value

        print(f"  ID: {af_id}")
        print(f"  Titel: {titel}")
        print(f"  Inhalt: {inhalt}")
        print(f"  Typ: {typ}")

        print(f"\nVerschiedene Werte:")
        print(f"  {os.path.basename(file1_path)}: {unique_values[0]}")
        print(f"  {os.path.basename(file2_path)}: {unique_values[1]}")

        if column_name == "Umsetzung":
            unique_values = "Ja", "Nein", "Teilweise"
        options = [str(val) for val in unique_values]
        options_dict = {str(i + 1): option for i, option in enumerate(options)}
        choice = multiple_choice(options_dict, f"\nWählen Sie einen Wert für '{column_name}':", multi = False)
        return options_dict[choice]

# Vereint zwei gleichartige Tabellen
def merge_tables(file1_path, file2_path):
    try:
        df1 = pd.read_excel(file1_path, sheet_name=0, skiprows=4)
        df2 = pd.read_excel(file2_path, sheet_name=0, skiprows=4)
        if len(df1) != len(df2):
            print("Verschiedene Tabellen. Vorgang wird abgebrochen.")
            exit(1)
        if id_check(file1_path) != id_check(file2_path):
            print("Eine Datei hat eindeutige, die andere nicht eindeutige IDs, Vorgang wird abgebrochen")
            exit(1)
    except PermissionError:
        print(f"Fehler: Eine der Dateien ist möglicherweise geöffnet und kann nicht gelesen werden. Vorgang wird abgebrochen.")
        exit(1)
    except Exception as e:
        print(f"Ein unerwarteter Fehler ist beim Öffnen der Datei aufgetreten: {e}")
        exit(1)

    columns_to_sum = ["Begründung für Entbehrlichkeit", "Verantwortlich", "Bemerkungen / Begründung für Nicht-Umsetzung", "Kostenschätzung"]

    # Hilfsfunktion zum Zusammenfassen von Feldern
    def concatenate_values(series):
        return ", ".join(map(str, series.dropna().unique()))

    agg_dict = {
        **{col: "first" for col in ["ID-Anforderung", "Titel", "Inhalt", "Typ"]},
        **{col: concatenate_values for col in columns_to_sum},
        "Umsetzung": lambda x: resolve_conflict(x, "Umsetzung", file1_path, file2_path),
        "Entbehrlich": lambda x: resolve_conflict(x, "Entbehrlich", file1_path, file2_path),
        "Umsetzung bis": lambda x: resolve_conflict(x, "Umsetzung bis", file1_path, file2_path)
    }
    merged_df = pd.concat([df1, df2]).groupby(["ID-Anforderung", "Titel", "Inhalt", "Typ"], sort=False, as_index=False).agg(agg_dict)

    merged_df = merged_df[
        ["ID-Anforderung", "Titel", "Inhalt", "Typ", "Entbehrlich", "Begründung für Entbehrlichkeit", "Umsetzung",
         "Umsetzung bis", "Verantwortlich", "Bemerkungen / Begründung für Nicht-Umsetzung", "Kostenschätzung"]]

    output_file_name = os.path.basename(file1_path)
    output_dir = os.path.join(os.path.dirname(file1_path), "Merged")
    os.makedirs(output_dir, exist_ok=True)
    output_file_path = os.path.join(output_dir, output_file_name)

    wb = openpyxl.load_workbook(file1_path)

    ws = wb.active
    for idx, row in merged_df.iterrows():
        excel_row = 6 + idx
        for col_name, value in row.items():
            col_letter = openpyxl.utils.get_column_letter(merged_df.columns.get_loc(col_name)+2)
            ws[f"{col_letter}{excel_row}"].value = value

    wb.save(output_file_path)

    print(f"\n{os.path.basename(file1_path)} und {os.path.basename(file2_path)} erfolgreich zusammengeführt und gespeichert unter: {output_file_path}")

# Behandelt die Modifizierung von allen Word-Dokumenten eines Ordners
def modify(directory):
    if not os.path.isdir(directory):
        print(f"Fehler: '{directory}' ist kein gültiges Verzeichnis.")
        return

    output_directory = os.path.join(directory, "Modified")
    os.makedirs(output_directory, exist_ok=True)

    section_titles = [
        "Beschreibung",
        "Gefährdungslage",
        "Anforderungen",
        "Basis-Anforderungen",
        "Standard-Anforderungen",
        "Anforderungen bei erhöhtem Schutzbedarf",
        "Weiterführende Informationen",
    ]

    anforderungen_subsections = [
        "Basis-Anforderungen",
        "Standard-Anforderungen",
        "Anforderungen bei erhöhtem Schutzbedarf",
    ]

    sections_to_remove_flags = {}
    anforderungen_selected = False

    print("Wählen Sie die Abschnitte aus, die entfernt werden sollen:")
    for title in section_titles:
        if title in anforderungen_subsections and anforderungen_selected:
            sections_to_remove_flags[title] = True
            continue

        sections_to_remove_flags[title] = False

        eingabe = input(f"- '{title}' entfernen? (ja/Nein): ")
        if eingabe.lower() in ["j", "ja"]:
            sections_to_remove_flags[title] = True
            if title == "Anforderungen":
                anforderungen_selected = True

    delete_entfallen = True
    integrate_checklists = False
    implemented = True
    partly_implemented = False
    not_implemented = False
    unnecessary = True
    checklist_path = None
    if not anforderungen_selected:
        delete_entfallen = ask_user("- Entfallene Anforderungen entfernen?")
        integrate_checklists = ask_user("- Excel Checklisten integrieren?", default_yes=False)
        if integrate_checklists:
            while True:
                checklist_path = input("\nPfad zum Ordner mit den Checklisten: ").strip().strip('"')
                if os.path.isdir(checklist_path):
                    break
                else:
                    print("Kein gültiger Pfad zu einem Ordner")
            print("\nHinweis: Es werden nur Anforderungen entfernt, bei denen alle Teil-Anforderungen die nachfolgenden Auswahlmöglichkeiten erfüllen.\n")
            implemented = ask_user("- In Tabellen bereits umgesetzte Anforderungen entfernen?")
            partly_implemented = ask_user("- In Tabellen teilweise umgesetzte Anforderungen entfernen?", default_yes=False)
            not_implemented = ask_user("- In Tabellen nicht umgesetzte Anforderungen entfernen?", default_yes=False)
            unnecessary = ask_user("- In Tabellen entbehrliche Anforderungen entfernen?")

    save_as_pdf = ask_user("\n- Dokumente auch als PDF speichern? Hinweis: Diese Operation ist sehr rechenaufwendig.", default_yes=False)

    sections_for_removal = [title for title, remove in sections_to_remove_flags.items() if remove]

    if not sections_for_removal and not delete_entfallen and not save_as_pdf and not integrate_checklists:
        print("Keine Aktion ausgewählt.")
        return

    for filename in os.listdir(directory):
        if is_docx(filename):
            original_file_path = os.path.join(directory, filename)
            output_file_path = os.path.join(output_directory, filename)

            dokument = Document(original_file_path)
            if sections_for_removal:
                remove_sections(dokument, sections_for_removal, section_titles)
            if delete_entfallen:
                remove_entfallene_anforderungen(dokument)
            if integrate_checklists:
                checklist_integration(checklist_path, filename, dokument, implemented, partly_implemented, not_implemented, unnecessary)
            try:
                dokument.save(output_file_path)
                if save_as_pdf:
                    pdf_directory = os.path.join(output_directory, "PDFs")
                    os.makedirs(pdf_directory, exist_ok=True)
                    pdf_filename = os.path.splitext(filename)[0] + ".pdf"
                    pdf_path = os.path.join(pdf_directory, pdf_filename)
                    docx2pdf.convert(output_file_path, pdf_path)
            except Exception as e:
                print(f" Fehler beim Speichern von {filename}: {e}")

    if not any(is_docx(f) for f in os.listdir(directory)):
        print("\nKeine .docx-Dateien im Verzeichnis gefunden.")
    else:
        print(f"\nDateien erfolgreich gespeichert.")

# Entfernt gewünschte Abschnitte eines Word-Dokuments
def remove_sections(dokument, sections_to_remove, section_titles):
    paragraphs = dokument.paragraphs
    num_paragraphs = len(paragraphs)

    found_sections_indices = {}
    for i, p in enumerate(paragraphs):
        p_text = p.text.strip()
        for title in sorted(section_titles, key=len, reverse=True): # Wichtig, damit 'Anforderungen' und 'Anforderungen des erhöhten Schutzbedarfs' nicht vertauscht werden
            if p_text.startswith(title) and title not in found_sections_indices:
                found_sections_indices[title] = i
                break

    section_ranges = {}
    sorted_found_keys = sorted(
        found_sections_indices.items(),
        key=lambda item: section_titles.index(item[0]) if item[0] in section_titles else float('inf')
    )
    sorted_found_keys.sort(key=lambda item: item[1])

    for i, (title, start_index) in enumerate(sorted_found_keys):
        end_index = num_paragraphs
        if i + 1 < len(sorted_found_keys):
            _, next_start_index = sorted_found_keys[i + 1]
            end_index = next_start_index
        if start_index < end_index:
            section_ranges[title] = {"start": start_index, "end": end_index}

    elements_to_delete = []
    paragraph_ranges_to_delete = [section_ranges[title] for title in sections_to_remove if title in section_ranges]
    paragraph_ranges_to_delete.sort(key=lambda r: r["start"], reverse=True)

    body_elements = list(dokument.element.body.iterchildren())
    for range_info in paragraph_ranges_to_delete:
        start_idx = range_info["start"]
        end_idx = range_info["end"]

        try:
            start_paragraph = paragraphs[start_idx]
            start_element = start_paragraph._element
            end_element = paragraphs[end_idx]._element if end_idx < num_paragraphs else None
        except IndexError:
            continue

        body_end_idx = len(body_elements)
        if end_element in body_elements:
            body_end_idx = body_elements.index(end_element)

        for i in range(body_end_idx - 1, -1, -1):
            element = body_elements[i]
            if element == start_element:
                if isinstance(element, CT_P):
                    elements_to_delete.append(element)
                break
            if isinstance(element, (CT_P, CT_Tbl)):
                elements_to_delete.append(element)

    for element in elements_to_delete:
        if element.getparent() is not None:
            element.getparent().remove(element)
    return

# Entfernt entfallene Anforderungen aus Docx-Dokument
def remove_entfallene_anforderungen(dokument):
    paragraphs = dokument.paragraphs
    body_elements = list(dokument.element.body.iterchildren())

    entfallen_pattern = re.compile(r"\b[A-Z]{2,}\.(\d+\.){1,3}A\d+\s+ENTFALLEN\b")
    anforderung_heading_pattern = re.compile(r"\b[A-Z]{2,}\.(\d+\.){1,3}A\d+\b")

    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        if entfallen_pattern.search(para.text):
            start_index = i
            i += 1
            while i < len(paragraphs):
                next_para_text = paragraphs[i].text.strip()
                if anforderung_heading_pattern.match(next_para_text) or next_para_text in ["Standard-Anforderungen", "Anforderungen bei erhöhtem Schutzbedarf", "Weiterführende Informationen"]:
                    break
                i += 1
            for j in range(start_index, i):
                el = paragraphs[j]._element
                if el in body_elements and el.getparent() is not None:
                    el.getparent().remove(el)
            continue
        i += 1

# Integriert die Werte der Checklisten ins Dokument, indem Anforderungen, die nicht die Bedingungen erfüllen, entfernt werden
def checklist_integration(checklist_path, filename, dokument, implemented, partly_implemented, not_implemented, unnecessary):
    docx_short = filename.split()[0]
    if docx_short.startswith("CCON"):
        docx_short = "CON" + docx_short[4:]

    found_checklist = None
    for checklist in os.listdir(checklist_path):
        if get_name(checklist, True) == docx_short:
            found_checklist = os.path.join(checklist_path, checklist)
    if found_checklist is None:
        print(f"Keine passende Checkliste für {filename} gefunden")
        return

    df = load_data(found_checklist)
    id_status = id_check(found_checklist)
    if id_status == 2:
        df["ID-Anforderung"] = df["ID-Anforderung"].astype(str).apply(lambda x: x[:-1] if len(x) > 1 else x)
    if id_status == 3:
        print(f"In {found_checklist} gibt es sowohl eindeutige als auch nicht eindeutige IDs, diese Datei wird übersprungen für die Integration")
        return
    df = remove_specific_requirements(df, "Titel", "ENTFALLEN")
    if implemented:
        df = remove_specific_requirements(df, "Umsetzung", "Ja")
    if partly_implemented:
        df = remove_specific_requirements(df, "Umsetzung", "Teilweise")
    if not_implemented:
        df2 = get_specific_df_with_two_conditions_and_emptiness(df, "Umsetzung", "Nein", "Entbehrlich", "Nein")
        contents_to_remove = set(df2["Inhalt"].unique())
        df = df[~df["Inhalt"].isin(contents_to_remove)] # ~not
    if unnecessary:
        df = remove_specific_requirements(df, "Entbehrlich", "Ja")

    df = df[["ID-Anforderung", "Inhalt"]]

    paragraphs = dokument.paragraphs
    body_elements = list(dokument.element.body.iterchildren())

    valid_ids = set(df["ID-Anforderung"].unique())

    # Falls keine Anforderung übrig ist, gesamten Abschnitt entfernen
    if not valid_ids:
        section_titles = [
            "Beschreibung",
            "Anforderungen",
            "Weiterführende Informationen",
        ]
        sections_to_remove_flags = {}
        sections_to_remove_flags["Anforderungen"] = True
        sections_for_removal = [title for title, remove in sections_to_remove_flags.items() if remove]
        remove_sections(dokument, sections_for_removal, section_titles)

    heading_pattern = re.compile(r"\b([A-Z]{2,}\.(\d+\.){1,3}A\d+)\b")

    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        heading_match = heading_pattern.match(para.text.strip())
        if heading_match:
            current_id = heading_match.group(1)
            start_index = i
            i += 1
            while i < len(paragraphs):
                next_para_text = paragraphs[i].text.strip()
                if (heading_pattern.match(next_para_text) or next_para_text in
                        ["Standard-Anforderungen", "Anforderungen bei erhöhtem Schutzbedarf", "Weiterführende Informationen"]):
                    break
                i += 1
            end_index = i

            if current_id not in valid_ids:
                for j in range(start_index, end_index):
                    el = paragraphs[j]._element
                    if el in body_elements and el.getparent() is not None:
                        el.getparent().remove(el)
        else:
            i += 1

# Speichert einen Dataframe in einem gewünschten Format
def save_df(df, export_file_path, index_number):
    file_format_choice_key = multiple_choice(mapping.format_options,"\nIn welchem Format soll die Tabelle gespeichert werden?", multi=False)

    if file_format_choice_key == "1":
        df.to_excel(export_file_path + ".xlsx", index=index_number)
    elif file_format_choice_key == "2":
        df.to_csv(export_file_path + ".csv", index=index_number)
    elif file_format_choice_key == "3":
        df.to_json(export_file_path + ".json", orient="records", indent=4, force_ascii=False)
    elif file_format_choice_key == "4":
        df.to_markdown(export_file_path + ".md", index=index_number)
    elif file_format_choice_key == "5":
        df.to_html(export_file_path + ".html", index=index_number)
    elif file_format_choice_key == "6":
        df.to_xml(export_file_path + ".xml", index=index_number)
    elif file_format_choice_key == "7":
        df.to_csv(export_file_path + ".txt", sep=" ", index=index_number)
    elif file_format_choice_key == "8":
        df.to_csv(export_file_path + ".tsv", sep="\t", index=index_number)
    elif file_format_choice_key == "9":
        delimiter = input("\nTrennzeichen: ")
        df.to_csv(export_file_path + ".txt", sep=delimiter, index=index_number)
    else:
        print("Keine gültige Auswahl, wird als .xlsx Datei gespeichert")
        df.to_excel(export_file_path + ".xlsx", index=index_number)

    print("Datei erfolgreich gespeichert")

# Durchsucht die Tabellen eines Ordners nach verschiedenen Möglichkeiten
def search(directory):
    if not os.path.isdir(directory):
        print(f"Fehler: '{directory}' ist kein gültiges Verzeichnis.")
        return

    all_dfs = []
    print("Lade Dateien...")
    for file_name in os.listdir(directory):
        file_path = os.path.join(directory, file_name)
        if is_format(file_name):
            df = load_data(file_path)
            if df is not None:
                all_dfs.append(df)

    if all_dfs:
        df = pd.concat(all_dfs, ignore_index=True)
    else:
        print("Keine Datei im Format Checkliste_XXX.X.X.xlsx gefunden")
        return
    df = df[["ID-Anforderung", "Titel", "Inhalt", "Typ", "Entbehrlich", "Begründung für Entbehrlichkeit", "Umsetzung", "Umsetzung bis", "Verantwortlich", "Bemerkungen / Begründung für Nicht-Umsetzung", "Kostenschätzung"]]

    choice = input("Auf welche Art sollen die Tabellen durchsucht werden?\n\n1 - Text\n2 - Regex\n3 - SQL\n\nAuswahl (1-3): ")
    if choice == "1":
        search_term = input("\nSuchbegriff: ")
        mask = df.apply(lambda row: row.astype(str).str.contains(search_term, case=False).any(), axis=1)
        df = df[mask]
    elif choice == "2":
        regex_pattern = input("\nRegulärer Ausdruck: ")
        mask = df.apply(lambda row: row.astype(str).str.contains(regex_pattern, regex=True).any(), axis=1)
        df = df[mask]
    elif choice == "3":
        sql_pattern = input("\nDie SQL-Abfrage muss \"FROM df\" beinhalten.\nEingabe: ")
        df = duckdb.query(sql_pattern).df()
    else: print("Ungültige Auswahl, Vorgang wird abgebrochen")

    if len(df) == 1:
        print("\n1 Ergebnis gefunden")
    else: print(f"\n{len(df)} Ergebnisse gefunden")
    if len(df) == 0:
        return

    export_dir = os.path.join(os.path.abspath(directory), "Search result")
    os.makedirs(export_dir, exist_ok=True)
    export_file_path = os.path.join(export_dir, f"Search result {datetime.today().strftime('%d-%m-%Y')}")
    save_df(df, export_file_path, False)

# Analysiert die abgedeckten elementaren Gefährdungen einer Datei
def risk_analysis(file_path):
    krt_map = {item['id']: item for item in mapping.krt}
    df = load_data(file_path)
    implemented = get_specific_df(df, "Umsetzung", "ja")
    covered_risks = set()
    covered_cia = set()

    id_status = id_check(file_path)
    if id_status == 1:
        ids = implemented["ID-Anforderung"].unique()
    elif id_status == 2:
        ids = [req_id[:-1] for req_id in implemented["ID-Anforderung"].unique()]
    else:
        print(
            f"Die Datei {file_path} enthält sowohl eindeutige als auch nicht eindeutige IDs, möglicherweise können nicht alle Gefährdungen zugeordnet werden")
        ids = implemented["ID-Anforderung"].unique()

    for req_id in ids:
        item = krt_map.get(req_id)
        covered_risks.update(item['gefahren'])
        if item['cia']:
            for char in item['cia']:
                covered_cia.update(char)

    return covered_risks, covered_cia

# Handler für Datei oder Ordner
def risk_handler(file_or_dir):
    if os.path.isfile(file_or_dir) and is_format(file_or_dir):
        covered_risks_codes, covered_cia_codes = risk_analysis(file_or_dir)

        print(f"\n--- Analyse für Datei: {os.path.basename(file_or_dir)} ---")

        if covered_risks_codes:
            print("(Teilweise) abgedeckte elementare Gefährdungen:")
            for risk_code in sorted(list(covered_risks_codes)):
                print(f"- {risk_code}: {mapping.gefahren[risk_code]}")
        else:
            print("\nKeine elementaren Gefährdungen abgedeckt.")

        if covered_cia_codes:
            print("\n(Teilweise) abgedeckte CIA-Attribute:")
            for cia_code in sorted(list(covered_cia_codes)):
                full_name = mapping.cia.get(cia_code)
                print(f"- {full_name} ({cia_code})")
        else:
            print("\nKeine CIA-Attribute abgedeckt.")

    elif os.path.isdir(file_or_dir):
        all_covered_risks_across_files = set()
        all_covered_cia_across_files = set()
        processed_files_count = 0

        for file_name in os.listdir(file_or_dir):
            file_path = os.path.join(file_or_dir, file_name)
            if os.path.isfile(file_path) and is_format(file_name):
                covered_risks_for_file, covered_cia_for_file = risk_analysis(file_path)

                all_covered_risks_across_files.update(covered_risks_for_file)
                all_covered_cia_across_files.update(covered_cia_for_file)
                processed_files_count += 1
            else:
                print(f"Überspringe: {file_name} (Kein gültiges Dateiformat)")

        if processed_files_count == 0:
            print("Keine gültigen Dateien im Ordner gefunden")
            return

        print(f"\n--- Analyse für Ordner: {os.path.basename(file_or_dir)} ---")

        total_possible_risks = set(mapping.gefahren.keys())
        missing_risks = total_possible_risks - all_covered_risks_across_files

        print("\nSummierte (teilweise) abgedeckte elementare Gefährdungen:")
        if all_covered_risks_across_files:
            for risk_code in sorted(list(all_covered_risks_across_files)):
                print(f"- {risk_code}: {mapping.gefahren[risk_code]}")
        else:
            print("Keine elementaren Gefährdungen abgedeckt.")

        if missing_risks:
            print("\nAn keiner Stelle abgedeckte elementare Gefährdungen:")
            for risk_code in sorted(list(missing_risks)):
                print(f"- {risk_code}: {mapping.gefahren[risk_code]}")
        else:
            print("\nAlle möglichen elementaren Gefährdungen sind an mind. einer Stelle abgedeckt!")

        total_possible_cia = set(mapping.cia.keys())
        missing_cia = total_possible_cia - all_covered_cia_across_files

        if all_covered_cia_across_files:
            print("\nSummierte (teilweise) abgedeckte CIA-Attribute:")
            for cia_code in sorted(list(all_covered_cia_across_files)):
                full_name = mapping.cia.get(cia_code, cia_code)
                print(f"- {full_name} ({cia_code})")

        if missing_cia:
            print("\nNoch nicht an mind. einer Stelle abgedeckte CIA-Attribute:")
            for cia_code in sorted(list(missing_cia)):
                full_name = mapping.cia.get(cia_code, cia_code)
                print(f"- {full_name} ({cia_code})")

    else:
        print("Kein gültiger Pfad: Bitte geben Sie eine gültige Datei im Format Checkliste_XXX.X.X.xlsx oder einen Ordner an.")

# Setzt die IDs einer Datei, eindeutig bei 1, zurück auf nicht eindeutig bei 2
def id_setter(file, mode):
    workbook = openpyxl.load_workbook(file)
    sheet = workbook.active

    if mode == 1:
        id_counts = defaultdict(int)
        for row in range(6, sheet.max_row + 1):
            cell = sheet.cell(row=row, column=2)
            original_value = str(cell.value) if cell.value else ""
            if original_value:
                letter_to_add = string.ascii_lowercase[id_counts[original_value]]
                cell.value = f"{original_value}{letter_to_add}"
                id_counts[original_value] += 1

    elif mode == 2:
        for row in range(6, sheet.max_row + 1):
            cell = sheet.cell(row=row, column=2)
            if cell.value and isinstance(cell.value, str) and cell.value[-1].isalpha():
                cell.value = cell.value[:-1]

    workbook.save(file)

# Behandelt für Dateien und Ordner, wie mit der Eindeutigmachung der IDs umgegangen werden soll
def id_handler(file_or_dir):
    if os.path.isfile(file_or_dir) and is_format(file_or_dir):
        mode = id_check(file_or_dir)
        if mode not in [1,2]:
            print("Sowohl eindeutige als auch nicht eindeutige IDs in der Datei vorhanden. Vorgang wird abgebrochen.")
            return
        id_setter(file_or_dir, mode)
        if mode == 1: print("Eindeutige IDs erfolgreich hinzugefügt")
        elif mode == 2: print("Eindeutige IDs erfolgreich wieder entfernt")
    elif os.path.isdir(file_or_dir):
        print("Starte ID-Modifikation...")
        mode = id_check(file_or_dir)
        if mode not in [1,2]:
            print("Sowohl eindeutige als auch nicht eindeutige IDs gefunden, wie soll verfahren werden?")
            choice = input("\n1. Alle IDs eindeutig machen\n2. Alle IDs auf nicht eindeutig zurücksetzen\n3. Liste ausgeben\n4. Abbruch\n\nAuswahl: ").strip()
            if choice in ['1','2']:
                mode = int(choice)
            elif choice == '3':
                for file, status in mode.items():
                    print(f"- {file}: {status}")
                print("\nWie soll verfahren werden?")
                choice2 = input("\n1. Alle IDs eindeutig machen\n2. Alle IDs auf nicht eindeutig zurücksetzen\n3. Abbruch\n\nAuswahl: ").strip()
                if choice2 in ['1','2']:
                    mode = int(choice2)
                elif choice2 == '3':
                    print("Vorgang wird abgebrochen")
                    return
                else:
                    print("Ungültige Auswahl, Vorgang wird abgebrochen.")
                    return
            elif choice == '4':
                print("Vorgang wird abgebrochen.")
                return
            else:
                print("Ungültige Auswahl, Vorgang wird abgebrochen.")
                return

        for file_name in os.listdir(file_or_dir):
            if is_format(file_name):
                file_path = os.path.join(file_or_dir, file_name)
                id_setter(file_path, mode)
        if mode == 1: print("Eindeutige IDs erfolgreich hinzugefügt")
        elif mode == 2: print("Eindeutige IDs erfolgreich wieder entfernt")
    else: print("Keine gültige Datei gefunden")


# Vergleicht zwei Dataframes für die Snapshot-Funktion
def compare_dataframes(df1, df2):
    df1 = remove_specific_requirements(df1.copy(), "Titel", "ENTFALLEN")
    df2 = remove_specific_requirements(df2.copy(), "Titel", "ENTFALLEN")

    if len(df1) != len(df2):
        print(f"\nUnterschiedliche Zahl an Anforderungen gefunden ({abs(len(df1) - len(df2))}), Analysen werden womöglich beeinträchtigt")

    # Erstellt eine schöne Zahl für die Differenz
    def zahl(zahl_snapshot, zahl_aktuell):
        diff = zahl_aktuell - zahl_snapshot
        if diff > 0: return f"+{diff}"
        elif diff < 0: return str(diff)
        else: return "+/- 0"

    df1['Baustein-ID'] = df1['ID-Anforderung'].apply(get_baustein_from_id)
    df2['Baustein-ID'] = df2['ID-Anforderung'].apply(get_baustein_from_id)
    df1['row_key'] = df1.apply(get_row_key, axis=1)
    df2['row_key'] = df2.apply(get_row_key, axis=1)

    all_baustein_ids = sorted(list(set(df1['Baustein-ID'].dropna().unique()).union(set(df2['Baustein-ID'].dropna().unique()))))

    total_modified_fields = 0
    baustein_diffs = set()
    baustein_summary = {}

    for baustein_id in all_baustein_ids:
        df1_baustein = df1[df1['Baustein-ID'] == baustein_id]
        df2_baustein = df2[df2['Baustein-ID'] == baustein_id]
        modified_counter = 0

        common_keys = set(df1_baustein['row_key']).intersection(set(df2_baustein['row_key']))

        for comp_key in common_keys:
            row1 = df1_baustein[df1_baustein['row_key'] == comp_key].iloc[0]
            row2 = df2_baustein[df2_baustein['row_key'] == comp_key].iloc[0]

            for col in df1.columns.drop(['Baustein-ID', 'row_key', 'ID-Anforderung', 'Titel', 'Inhalt', 'Typ']):
                val1 = row1.get(col)
                val2 = row2.get(col)
                val1_str = str(val1).strip() if not pd.isna(val1) else ''
                val2_str = str(val2).strip() if not pd.isna(val2) else ''

                if val1_str != val2_str:
                    modified_counter += 1

        if modified_counter > 0:
            total_modified_fields += modified_counter
            baustein_diffs.add(baustein_id)

            _, _, fully_implemented1, _ = implementation_count(df1_baustein)
            _, _, fully_implemented2, _ = implementation_count(df2_baustein)
            text = "modifiziertes Feld" if modified_counter == 1 else "modifizierte Felder"
            baustein_summary[baustein_id] = f"{baustein_id}: Umgesetzte Anforderungen: {zahl(fully_implemented1, fully_implemented2)} ({modified_counter} {text})"

    not_implemented1, partly_implemented1, fully_implemented1, unnecessary1 = implementation_count(df1)
    not_implemented2, partly_implemented2, fully_implemented2, unnecessary2 = implementation_count(df2)
    cost1, cost2 = cost_count(df1), cost_count(df2)
    all_responsibles1, all_responsibles2 = f"{', '.join(get_responsible(df1))}", f"{', '.join(get_responsible(df2))}"

    ni1 = get_specific_df_with_two_conditions_and_emptiness(df1, "Umsetzung", "nein", "Entbehrlich", "nein")
    ni2 = get_specific_df_with_two_conditions_and_emptiness(df2, "Umsetzung", "nein", "Entbehrlich", "nein")
    ni_basis1, ni_standard1, ni_high1 = get_type(ni1)
    ni_basis2, ni_standard2, ni_high2 = get_type(ni2)

    print("\n ------ Übersicht ------\n")
    print(f"Modifizierte Felder: {total_modified_fields}")
    print(f"Umgesetzte Anforderungen: {zahl(fully_implemented1, fully_implemented2)}")
    print(f"Teilweise umgesetzte Anforderungen: {zahl(partly_implemented1, partly_implemented2)}")
    print(f"Entbehrliche Anforderungen: {zahl(unnecessary1, unnecessary2)}")
    print(f"Nicht umgesetzte Anforderungen: {zahl(not_implemented1, not_implemented2)}")
    print(f"  (Davon Basis: {zahl(ni_basis1, ni_basis2)}, Standard: {zahl(ni_standard1, ni_standard2)}, Erhöhter Schutzbedarf: {zahl(ni_high1, ni_high2)})")
    print(f"Summierte Kostenschätzung: {zahl(cost1, cost2)}€")
    print(f"Verantwortliche Snapshot: {all_responsibles1}")
    print(f"Verantwortliche Aktuell: {all_responsibles2}")

    snapshot_ids = set(df1['Baustein-ID'].dropna().unique())
    current_ids = set(df2['Baustein-ID'].dropna().unique())

    added_ids = sorted(list(current_ids - snapshot_ids))
    removed_ids = sorted(list(snapshot_ids - current_ids))

    if added_ids or removed_ids:
        for baustein_id in added_ids:
            print(f"\nBaustein hinzugefügt: {baustein_id}")
        for baustein_id in removed_ids:
            print(f"\nBaustein entfernt: {baustein_id}")

    if not baustein_diffs:
        return

    if ask_user("\nÄnderungen pro Baustein anzeigen?"):
        print("\n ------ Änderungen pro Baustein ------\n")
        for baustein_id in all_baustein_ids:
            if baustein_id in baustein_summary:
                print(baustein_summary[baustein_id])

    detail_options = {
        "1": "Alle Differenzen anzeigen",
        "2": "Spezifischen Baustein auswählen",
        "3": "Beenden"
    }
    detail_choice = multiple_choice(detail_options, "\nDetaillierte Ansicht der Änderungen:", multi=False, default_value="3")

    if detail_choice == "1":
        for baustein_id in sorted(list(baustein_diffs)):
            print(f"\n--- Änderungen in Baustein {baustein_id} ---")
            display_detailed_baustein_diff(df1[df1['Baustein-ID'] == baustein_id], df2[df2['Baustein-ID'] == baustein_id])
    elif detail_choice == "2":
        baustein_selection_options = {str(i + 1): bid for i, bid in enumerate(sorted(list(baustein_diffs)))}
        selected_baustein_idx = multiple_choice(baustein_selection_options, "\nWählen Sie einen Baustein:", multi=False)
        selected_baustein_id = baustein_selection_options[selected_baustein_idx]
        print(f"\n--- Änderungen in Baustein {selected_baustein_id} ---")
        display_detailed_baustein_diff(df1[df1['Baustein-ID'] == selected_baustein_id], df2[df2['Baustein-ID'] == selected_baustein_id])
    else:
        print("Vorgang beendet.")

# Zeigt die Differenzen zwischen Baustein aus Snapshot und aktuellem Stand an
def display_detailed_baustein_diff(df1_baustein, df2_baustein):
    all_keys_1 = set(df1_baustein['row_key'])
    all_keys_2 = set(df2_baustein['row_key'])

    common_keys = all_keys_1.intersection(all_keys_2)

    for comp_key in common_keys:
        row1 = df1_baustein[df1_baustein['row_key'] == comp_key].iloc[0]
        row2 = df2_baustein[df2_baustein['row_key'] == comp_key].iloc[0]

        for col in df1_baustein.columns.drop(['Baustein-ID', 'row_key']):
            val1 = row1.get(col)
            val2 = row2.get(col)

            val1_str = str(val1).strip() if not pd.isna(val1) else ''
            val2_str = str(val2).strip() if not pd.isna(val2) else ''

            if val1_str != val2_str:
                print(f"  ID-Anforderung: {row1.get('ID-Anforderung')}, Titel: {row1.get('Titel')}")
                print(f"    Feld: {col}")
                print(f"      Snapshot: {val1_str}")
                print(f"      Aktuell:  {val2_str}")

# Füllt eine Kopie der Checklisten mit den Daten aus Snapshot
def restore_snapshot(directory, snapshot_df):
    restored_dir = os.path.join(directory, "Restored")
    if not os.path.exists(restored_dir):
        os.makedirs(restored_dir)

    excel_files_to_restore = []
    for file_name in os.listdir(directory):
        file_path = os.path.join(directory, file_name)
        if os.path.isfile(file_path) and is_format(file_name):
            destination_path = os.path.join(restored_dir, file_name)
            shutil.copy2(file_path, destination_path)
            excel_files_to_restore.append(destination_path)

    if not excel_files_to_restore:
        print("Keine passenden Excel-Dateien für die Wiederherstellung gefunden.")
        return

    snapshot_df['row_key'] = snapshot_df.apply(get_row_key, axis=1)

    relevant_cols = [
        "Entbehrlich", "Begründung für Entbehrlichkeit", "Umsetzung",
        "Umsetzung bis", "Verantwortlich",
        "Bemerkungen / Begründung für Nicht-Umsetzung", "Kostenschätzung"
    ]
    key_cols = ["ID-Anforderung", "Inhalt"]

    for file_path in excel_files_to_restore:
        try:
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active

            header_row_index = 5

            excel_header_map = {
                cell.value: cell.column
                for cell in sheet[header_row_index] if cell.value is not None
            }

            id_anforderung_col_excel = excel_header_map.get(key_cols[0])
            inhalt_col_excel = excel_header_map.get(key_cols[1])

            col_indices_to_update = {
                col: excel_header_map[col]
                for col in relevant_cols if col in excel_header_map
            }

            for excel_row_idx in range(header_row_index + 1, sheet.max_row + 1):
                id_anforderung_val = sheet.cell(row=excel_row_idx, column=id_anforderung_col_excel).value
                inhalt_val = sheet.cell(row=excel_row_idx, column=inhalt_col_excel).value

                if id_anforderung_val is None or inhalt_val is None:
                    continue
                current_row_key = f"{id_anforderung_val}_{inhalt_val}"

                matching_snapshot_rows = snapshot_df[snapshot_df['row_key'] == current_row_key]

                if not matching_snapshot_rows.empty:
                    snapshot_row = matching_snapshot_rows.iloc[0]
                    for col_name, excel_col_idx in col_indices_to_update.items():
                        sheet.cell(row=excel_row_idx, column=excel_col_idx).value = snapshot_row.get(col_name)
                else:
                    print(f"Hinweis: Zeile mit ID {id_anforderung_val}' in '{os.path.basename(file_path)}' wurde nicht im Snapshot gefunden.")

            workbook.save(file_path)

        except Exception as e:
            print(f"Fehler beim Wiederherstellen von '{os.path.basename(file_path)}': {e}")

    print("\nSnapshot-Wiederherstellung abgeschlossen.")

# Erstellt, vergleicht oder stellt einen Snapshot der Tabellen wieder her
def snapshot(directory):
    if not os.path.isdir(directory):
        print(f"Fehler: '{directory}' ist kein gültiges Verzeichnis.")
        return

    all_dfs = []
    print("Lade Dateien...")
    for file_name in os.listdir(directory):
        file_path = os.path.join(directory, file_name)
        if is_format(file_name):
            df = load_data(file_path)
            if df is not None:
                all_dfs.append(df)

    if all_dfs:
        df = pd.concat(all_dfs, ignore_index=True)
        id_status = id_check(directory)
        if id_status == 2:
            df["ID-Anforderung"] = df["ID-Anforderung"].astype(str).apply(lambda x: x[:-1] if len(x) > 1 else x)
        if id_status == 3:
            print("\nHinweis: Sowohl eindeutige als auch nicht eindeutige IDs gefunden, führt eventuell zu Problemen mit Snapshots")
    else:
        print("Keine Datei im Format Checkliste_XXX.X.X.xlsx gefunden")
        return
    df = df[["ID-Anforderung", "Titel", "Inhalt", "Typ", "Entbehrlich", "Begründung für Entbehrlichkeit", "Umsetzung",
             "Umsetzung bis", "Verantwortlich", "Bemerkungen / Begründung für Nicht-Umsetzung", "Kostenschätzung"]]

    snapshots_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Snapshots")
    snapshot_files = glob.glob(os.path.join(snapshots_dir, "Snapshot_*.csv"))

    options = {
        "1": "Snapshot erstellen",
        "2": "Snapshot mit aktuellem Stand vergleichen",
        "3": "Snapshot wiederherstellen",
        "4": "Abbrechen"
    }
    choice = multiple_choice(options, "\nWählen Sie eine Option aus:", multi = False, default_value="4")
    if choice == "1":
        if not os.path.exists(snapshots_dir):
            os.makedirs(snapshots_dir)

        snapshot_file_name = f"Snapshot_{datetime.today().strftime('%Y-%m-%d_%H-%M')}.csv"
        snapshot_file_path = os.path.join(snapshots_dir, snapshot_file_name)

        df.to_csv(snapshot_file_path, index=False)
        print(f"Snapshot erfolgreich gespeichert.")

    elif choice == "2" or choice == "3":
        if not snapshot_files:
            print("Keine Snapshots gefunden.")
            return

        snapshot_options = {str(i + 1): os.path.basename(f) for i, f in enumerate(sorted(snapshot_files))}

        snapshot_choice = multiple_choice(snapshot_options, "\nWählen Sie einen Snapshot aus:", multi=False)

        selected_snapshot_file = os.path.join(snapshots_dir, snapshot_options[snapshot_choice])
        snapshot_df = pd.read_csv(selected_snapshot_file)
        if choice == "2":
            print("\nFühre Vergleich durch...")
            compare_dataframes(snapshot_df, df)
        if choice == "3":
            print(f"\nStelle Snapshot wieder her...")
            restore_snapshot(directory, snapshot_df)
    else:
        print("Vorgang wird abgebrochen")

# Importiert Werte aus Dateien in verschiedenen Formaten
def import_files(checklist_directory):
    if not os.path.isdir(checklist_directory):
        print(f"Fehler: '{checklist_directory}' ist kein gültiges Verzeichnis.")
        return
    print(f"\nVerzeichnis der Checklisten: {checklist_directory}")
    files_list = []
    while True:
        import_dir_or_file = input("\nPfad zu Ordner oder Datei für Import: ").strip().strip('"')
        if os.path.isdir(import_dir_or_file):
            for file in os.listdir(import_dir_or_file):
                path = os.path.join(import_dir_or_file, file)
                files_list.append(path)
            break
        elif os.path.isfile(import_dir_or_file):
            files_list.append(import_dir_or_file)
            break
        else: print("Kein gültiger Pfad")
    if id_check(checklist_directory) != 1:
        print("\nModifzierte IDs in den Checklisten gefunden, beinträchtigt möglicherweise die Zuordnung beim Import.")

    formats = {
        "1": "Excel",
        "2": "CSV",
        "3": "JSON",
        "4": "Leerzeichengetrennt",
        "5": "Tabstoppgetrennt",
        "6": "Benutzerdefiniertes Trennzeichen",
    }
    file_format = multiple_choice(formats, "\nIn welchem Format sind die zu importierenden Dateien?", multi=False)
    if file_format == "1":
        skip = int(input("\nIn welcher Zeile liegen die Spaltenüberschriften? Auswahl (z.B. 1 oder 5): ")) - 1
    elif file_format == "6":
        delimiter = input("\nTrennzeichen: ")
        file_format = f"delimited:{delimiter}"

    dfs_to_merge = []

    for f_path in files_list:
        df = None
        if file_format == "1" and f_path.endswith('.xlsx'):
            df = pd.read_excel(f_path, skiprows=skip)
        elif file_format == "2" and f_path.endswith('.csv'):
            df = pd.read_csv(f_path)
        elif file_format == "3" and f_path.endswith('.json'):
            try:
                df = pd.read_json(f_path, orient='records')
            except ValueError as e:
                if "Trailing data" in str(e):
                    df = pd.read_json(f_path, lines=True, orient='records')
        elif file_format in ["4", "5"] or file_format.startswith("delimited:"):
            delimiter = " " if file_format == "4" else "\t" if file_format == "5" else \
                file_format.split(":")[1]
            if f_path.endswith(('.txt', '.tsv')):
                df = pd.read_csv(f_path, sep=delimiter)
        if df is not None and not df.empty:
            dfs_to_merge.append(df)

    merged_df = pd.concat(dfs_to_merge, ignore_index=True)

    import_columns = list(merged_df.columns)

    key_cols = ["ID-Anforderung", "Inhalt"]
    for key_col in key_cols:
        if key_col not in merged_df.columns:
            print(f"\nFehler: Die erforderliche Schlüsselspalte '{key_col}' wurde in den Importdaten nicht gefunden.")
            return
    non_import_cols = ["ID-Anforderung", "Titel", "Inhalt", "Typ"]

    relevant_cols = [
        "Entbehrlich", "Begründung für Entbehrlichkeit", "Umsetzung",
        "Umsetzung bis", "Verantwortlich",
        "Bemerkungen / Begründung für Nicht-Umsetzung", "Kostenschätzung"
    ]

    selectable_cols = [col for col in import_columns if col not in non_import_cols and col in relevant_cols]
    col_options = {str(i + 1): col for i, col in enumerate(selectable_cols)}

    if not col_options:
        print("\nKeine importierbaren Spalten gefunden, Vorgang wird abgebrochen.")
        return

    selected_keys = multiple_choice(
        col_options,
        "\nWelche dieser gefundenen Spalten sollen importiert werden? Enter für Alle",
        multi=True,
        default_value=list(col_options.keys())
    )
    relevant_cols = [col_options[key] for key in selected_keys]

    import_modes = {
        "1": "Vorhandene Daten überschreiben",
        "2": "Nur in leere Felder einfügen",
        "3": "Alle Daten zusammenfügen (mit Konfliktbehandlung)"
    }
    import_mode = multiple_choice(import_modes, "\nWählen Sie den Import-Modus:", multi=False)
    conflict_strategy = {}

    if import_mode == "1":
        print("\nHinweis: Daten, die nicht im Import vorkommen oder dort leer sind, werden in den Tabellen wie bisher beibehalten")
    elif import_mode == "3":
        conflict_cols = ["Umsetzung", "Umsetzung bis", "Entbehrlich"]
        print("\nLegen Sie die globale Strategie zur Konfliktlösung fest:")
        for col in conflict_cols:
            if col in relevant_cols:
                if col == "Umsetzung bis":
                    strategy_choice = multiple_choice({
                        "1": "Bisherige Daten aus Tabelle behalten",
                        "2": "Neue Daten aus Import übernehmen",
                        "3": "Beide Werte hinterlegen"
                    }, f"\nWenn bei '{col}' ein Konflikt besteht?", multi=False)
                    if strategy_choice == "2":
                        conflict_strategy[col] = "import"
                    elif strategy_choice == "3":
                        conflict_strategy[col] = "both"
                    else:
                        conflict_strategy[col] = "existing"
                else:
                    strategy_choice = multiple_choice({
                        "1": "Bisherige Daten aus Tabelle behalten",
                        "2": "Neue Daten aus Import übernehmen"
                    }, f"\nWenn bei '{col}' ein Konflikt besteht:", multi=False)
                    conflict_strategy[col] = "import" if strategy_choice == "2" else "existing"
        if "Kostenschätzung" in relevant_cols:
            conflict_strategy["Kostenschätzung"] = multiple_choice({
                "1": "Bisherige Daten aus Tabelle behalten",
                "2": "Neue Daten aus Import übernehmen",
                "3": "Kosten summieren",
                "4": "Mittelwert bilden",
                "5": "Beide Werte hinterlegen"
            }, f"\nWie soll mit 'Kostenschätzung' umgegangen werden?", multi=False)
        print("\nHinweis: Falls in zwei Feldern dieselben Werte vorkommen, werden sie nicht doppelt hinterlegt.")

    import_dir = os.path.join(checklist_directory, "Import")
    if not os.path.exists(import_dir):
        os.makedirs(import_dir)

    baustein_ids_from_import = set(merged_df['ID-Anforderung'].apply(get_baustein_from_id).dropna().unique())

    excel_files_to_import_to = []
    for file_name in os.listdir(checklist_directory):
        file_path = os.path.join(checklist_directory, file_name)
        if os.path.isfile(file_path) and is_format(file_name):
            checklist_baustein_id = get_name(file_name, True)
            if checklist_baustein_id in baustein_ids_from_import:
                destination_path = os.path.join(import_dir, file_name)
                shutil.copy2(file_path, destination_path)
                excel_files_to_import_to.append(destination_path)

    if not excel_files_to_import_to:
        print("Keine passenden Excel-Dateien zum Importieren gefunden.")
        return

    merged_df['row_key'] = merged_df.apply(get_row_key, axis=1)

    print("\nImportiere...")

    for file_path in excel_files_to_import_to:
        try:
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active
            header_row_index = 5

            excel_header_map = {cell.value: cell.column for cell in sheet[header_row_index] if cell.value is not None}
            id_anforderung_col_excel = excel_header_map.get(key_cols[0])
            inhalt_col_excel = excel_header_map.get(key_cols[1])

            col_indices_to_update = {col: excel_header_map.get(col) for col in relevant_cols if col in excel_header_map}

            for excel_row_idx in range(header_row_index + 1, sheet.max_row + 1):
                id_anforderung_val = sheet.cell(row=excel_row_idx, column=id_anforderung_col_excel).value
                inhalt_val = sheet.cell(row=excel_row_idx, column=inhalt_col_excel).value

                current_row_key = f"{id_anforderung_val}_{inhalt_val}"

                matching_merged_rows = merged_df[merged_df['row_key'] == current_row_key]

                if not matching_merged_rows.empty:
                    merged_row = matching_merged_rows.iloc[0]
                    for col_name, excel_col_idx in col_indices_to_update.items():
                        if not excel_col_idx: continue

                        import_value = merged_row.get(col_name)
                        if pd.isna(import_value):
                            import_value = None

                        current_cell = sheet.cell(row=excel_row_idx, column=excel_col_idx)

                        if import_mode == "1":
                            current_cell.value = import_value

                        elif import_mode == "2":
                            if current_cell.value is None or str(current_cell.value).strip() == "":
                                current_cell.value = import_value

                        elif import_mode == "3":
                            is_empty = current_cell.value is None or str(current_cell.value).strip() == ""

                            if is_empty:
                                current_cell.value = import_value
                            elif col_name == "Umsetzung bis" and col_name in conflict_strategy:
                                strategy = conflict_strategy[col_name]
                                if str(current_cell.value) != str(import_value):
                                    if strategy == 'import':
                                        current_cell.value = format_dates(import_value)
                                    elif strategy == 'both':
                                        current_cell.value = f"{str(format_dates(current_cell.value))}, {str(format_dates(import_value))}"
                            elif col_name == "Kostenschätzung" and col_name in conflict_strategy:
                                try:
                                    current_cost = float(current_cell.value) if current_cell.value is not None else 0.0
                                    imported_cost = float(import_value) if import_value is not None else 0.0

                                    cost_strat = conflict_strategy[col_name]
                                    if cost_strat == "2":
                                        current_cell.value = imported_cost
                                    elif cost_strat == "3":
                                        current_cell.value = current_cost + imported_cost
                                    elif cost_strat == "4":
                                        current_cell.value = (current_cost + imported_cost) / 2
                                    elif cost_strat == "5":
                                        if current_cost != imported_cost:
                                            current_cell.value = f"{current_cost}, {imported_cost}"
                                except (ValueError, TypeError):
                                    if str(current_cell.value) != str(import_value):
                                        current_cell.value = f"{str(current_cell.value)}, {str(import_value)}"
                            elif col_name in conflict_strategy:
                                if str(current_cell.value) != str(import_value):
                                    if conflict_strategy[col_name] == 'import':
                                        current_cell.value = import_value
                            else:
                                if str(current_cell.value) != str(import_value):
                                    current_cell.value = f"{str(current_cell.value)}, {str(import_value)}"

            workbook.save(file_path)
        except Exception as e:
            print(f"Fehler beim Importieren in '{os.path.basename(file_path)}': {e}")

    print("\nImport abgeschlossen.")

def main():
    parser = argparse.ArgumentParser(description='Arbeitet mit IT-Grundschutz-Check Excel Tabellen')
    parser.add_argument('file_or_dir', nargs="?", default=DEFAULT_PATH,
                        help='Pfad zur Datei oder zum Ordner (Kann im Code hinterlegt werden)')
    parser.add_argument('--mapping', action='store_true', help='Ordne der/den Tabelle(n) den Namen des Bausteins zu')
    parser.add_argument('--prozent', action='store_true', help='Zeigt wie viel Prozent eines Bausteins umgesetzt sind (ENTFALLEN und Entbehrlich wird ignoriert, Teilweise zählt als nicht umgesetzt)')
    parser.add_argument('--profil', action='store_true', help='Verschiebt nicht relevante Bausteine in einen Ordner nach IT-Grundschutz-Profilen (Menü öffnet sich); Auch für IT-Grundschutz-PDFs nutzbar')
    parser.add_argument('--ignore-standard', action='store_true', help='Setzt alle leeren Standard Anforderungen auf entbehrlich')
    parser.add_argument('--ignore-hoch', action='store_true', help='Setzt alle leeren Anforderungen des erhöhten Schutzbedarfs auf entbehrlich')
    parser.add_argument('--fully', action='store_true', help='(Mit --ignore-[...]) Setzt auch Anforderungen mit Umsetzung = Ja/Nein/Teilweise auf entbehrlich')
    parser.add_argument('--list', action='store_true', help='Listet Dateien eines Ordners nach verschiedene Kriterien (Menü öffnet sich)')
    parser.add_argument('--wiba-transfer', action='store_true', help='Markiert alle leeren, in WiBA enthaltenen Anforderungen als umgesetzt')
    parser.add_argument('--set-scale', action='store_true', help='Ändert die Umsetzungsskala. Führt evtl. zu Funktionseinschränkungen des Tools.')
    parser.add_argument('--export', action='store_true', help='Exportiere beliebige Spalten der Dateien in verschiedene Formate')
    parser.add_argument('--import-files', action='store_true', help='Importiert beliebige Spalten aus verschiedenen Dateiformaten')
    parser.add_argument('--report', action='store_true', help='Erstelle einen PDF-Report für eine oder alle Dateien')
    parser.add_argument('--merge', nargs=2, metavar=("file1", "file2"), help="Führt zwei Tabellen der selben Art zusammen und behandelt Konflikte")
    parser.add_argument('--modify', action='store_true', help='Modifiziert alle Bausteine eines Ordners im docx-Format, wahlweise Export zu PDF (Menü öffnet sich)')
    parser.add_argument('--search', action='store_true', help='Durchsuche alle Tabellen eines Ordners auf verschiedene Arten')
    parser.add_argument('--risks', action='store_true', help='Zeigt die von umgesetzten Anforderungen abgedeckten elementaren Gefährdungen an')
    parser.add_argument('--id', action='store_true', help='Gibt den einzelnen Anforderungen eine eindeutige ID bzw. entfernt sie wieder')
    parser.add_argument('--snapshot', action='store_true', help='Erstellt einen Snapshot des aktuellen Stands, vergleicht ihn oder stellt ihn wieder her')

    args = parser.parse_args()

    actions = {
        args.mapping: lambda: map_xlsx_files(args.file_or_dir),
        args.prozent: lambda: percentages(args.file_or_dir),
        args.profil: lambda: profile_handler(args.file_or_dir),
        args.ignore_standard: lambda: ignore_empty_requirements_handler(args.file_or_dir, "Typ", "Standard","Umsetzung", args.fully),
        args.ignore_hoch: lambda: ignore_empty_requirements_handler(args.file_or_dir, "Typ", "Hoch", "Umsetzung", args.fully),
        args.list: lambda: sort_list(args.file_or_dir),
        args.wiba_transfer: lambda: wiba_transfer(args.file_or_dir),
        args.set_scale: lambda: scale_setter_handler(args.file_or_dir),
        args.export: lambda: export_handler(args.file_or_dir),
        args.import_files: lambda: import_files(args.file_or_dir),
        args.report: lambda: report(args.file_or_dir),
        args.modify: lambda: modify(args.file_or_dir),
        args.search: lambda: search(args.file_or_dir),
        args.risks: lambda: risk_handler(args.file_or_dir),
        args.id: lambda: id_handler(args.file_or_dir),
        args.snapshot: lambda: snapshot(args.file_or_dir),
    }

    for condition, action in actions.items():
        if condition:
            action()
            return

    if args.merge:
        merge_tables(args.merge[0], args.merge[1])
        return

    if args.file_or_dir:
        if os.path.isfile(args.file_or_dir):
            analyze_single_file(args.file_or_dir)
        elif os.path.isdir(args.file_or_dir):
            analyze_all_files(args.file_or_dir)
        else:
            print("Datei oder Ordner nicht gefunden")
    else: print("Fehlendes Argument file_or_dir, kann übergeben oder im Code hinterlegt werden")

if __name__ == '__main__':
    main()
